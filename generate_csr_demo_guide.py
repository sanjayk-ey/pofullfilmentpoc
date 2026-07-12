"""
generate_csr_demo_guide.py
Builds the demo guide Word document that explains the four demo PO files and,
in particular, every CSR (human-in-the-loop) approval the AI raises for the
CSR-Approval-PO file.

Output: CSR_Demo_Guide.docx

Run:  python generate_csr_demo_guide.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
BLUE = RGBColor(0x25, 0x63, 0xEB)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUT = os.path.join(os.path.dirname(__file__), "CSR_Demo_Guide.docx")


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    p.space_after = Pt(4)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def body(doc, text, size=10.5, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        _set_cell(c, hd, bold=True, color=WHITE, size=9.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, "2563EB")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            _set_cell(cells[ci], val, size=9.5)
            if ri % 2 == 1:
                _shade(cells[ci], "EEF3FB")
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    return t


# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ── Cover ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Autonomous PO-to-Fulfillment")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = NAVY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Demo Files & CSR Approval Guide")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("How the Order assistant reads each purchase order and where it pauses "
                 "for Customer Service Rep (CSR) approval")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY

doc.add_paragraph()

# ── 1. The demo files ─────────────────────────────────────────────────────────
h1(doc, "1.  The four demo files")
body(doc, "There are four ready-to-use demo purchase orders in the demo/ folder — "
          "two plain-text (.txt) and two Excel (.xlsx). Each format has one "
          "happy-flow file and one CSR-approval file, so you can demo with either "
          "file type. The .txt and .xlsx twins contain identical data and behave "
          "exactly the same in the tool.")
table(doc,
      ["File", "Format", "Purpose", "PO Number", "Customer"],
      [
          ["Happy-Flow-PO.txt", "Text", "Happy flow / straight-through auto decision", "PO-2026-30001", "Great Lakes Plumbing Supply Co"],
          ["Happy-Flow-PO.xlsx", "Excel", "Happy flow / straight-through auto decision", "PO-2026-30001", "Great Lakes Plumbing Supply Co"],
          ["CSR-Approval-PO.txt", "Text", "Walks every CSR approval gate", "PO-2026-30002", "Great Lakes Plumbing Supply Co"],
          ["CSR-Approval-PO.xlsx", "Excel", "Walks every CSR approval gate", "PO-2026-30002", "Great Lakes Plumbing Supply Co"],
      ],
      widths=[2.0, 0.8, 2.5, 1.2, 2.0])
body(doc, "Tip: the two files are intentionally 'standard US PO' shaped — a header "
          "block (PO number, date, buyer company + email, ship-to, requested "
          "delivery date) followed by an order-lines table (Item #, Product Code, "
          "Description, Qty). Only PO Number, PO Date, buyer company + email, "
          "ship-to, requested delivery date, and per-line SKU/description/qty are "
          "needed — every other value (UOM, unit price, contact, contract, tax, "
          "freight) is looked up by the AI in the background from master data.",
     italic=True, color=GREY)

# ── 2. Running a demo ─────────────────────────────────────────────────────────
h1(doc, "2.  How to run a demo")
bullet(doc, "Start the app, then either paste the .txt contents into the PO box or "
            "upload the .xlsx file.")
bullet(doc, "The AI shows its work as it goes — it 'thinks' step-by-step (reading "
            "the PO, checking master data, taking decisions) and moves deliberately "
            "so the audience can follow. The screen auto-scrolls to follow the flow.")
bullet(doc, "When the AI needs a human decision it PAUSES and shows a decision card "
            "with buttons. Nothing proceeds until the CSR acts.")
bullet(doc, "To process a second PO, clear the conversation first so the new run "
            "starts clean.")

# ── 3. Happy flow ─────────────────────────────────────────────────────────────
h1(doc, "3.  Happy-Flow-PO  —  straight-through, no CSR needed")
body(doc, "This PO is clean: a known buyer, a complete ship-to, and three in-catalog "
          "products with valid quantities. Use it to show the fully autonomous, "
          "no-touch path.")
body(doc, "What the AI does automatically (no CSR pause):")
table(doc,
      ["Line", "Product Code", "Description", "Qty"],
      [
          ["1", "SKU-CTG-4520", "Ceramic Disc Faucet Cartridge", "100"],
          ["2", "SKU-SEL-1150", "Tank-to-Bowl Gasket Kit", "120"],
          ["3", "SKU-VLV-2201", "Pressure-Balancing Shower Valve", "15"],
      ],
      widths=[0.6, 1.6, 3.6, 0.8])
bullet(doc, "Intake: extracts all fields; backfills contact person, contract "
            "reference and other optional header fields from master data.")
bullet(doc, "Customer Validation: resolves the account hierarchy, customer tier, "
            "payment terms, distributor status and buying history.")
bullet(doc, "Product Match, Pricing & Promo, Approval, Credit, Inventory Checks, "
            "Logistics, Optimization: all pass within policy.")
bullet(doc, "Result: order auto-created. Final order total ≈ $3,348.86, customer "
            "confirmation issued with price, ETA, fulfillment source and tracking.")

# ── 4. CSR file: the planted issues ──────────────────────────────────────────
doc.add_page_break()
h1(doc, "4.  CSR-Approval-PO  —  what is 'wrong' in the file and why")
body(doc, "This single PO is deliberately built so the AI has to ask the CSR for a "
          "decision at every human-in-the-loop gate — first the seven intake / "
          "product-match gates, then four decision-layer gates (Pricing & Promo, "
          "Credit, Inventory, Logistics). That is every CSR approval process in the "
          "orchestration, all in one file. (Duplicate PO is the only case that needs "
          "no CSR approval — see section 7.) The table below explains what is planted "
          "on each part of the PO and the CSR gate it triggers.")

table(doc,
      ["On the PO", "What the AI detects", "CSR gate", "What the AI asks / shows", "CSR options"],
      [
          ["Buyer email: procurement@greatlakesps.com",
           "Email is not in the buyer directory",
           "Unresolved buyer",
           "Lists the buyers registered for this customer (e.g. John Miller, Linda Park)",
           "Pick a buyer / Type a name / Escalate"],
          ["Line 1: SKU-CTG-1000 (qty 40)",
           "SKU is obsolete / discontinued",
           "Obsolete-SKU substitution",
           "Recommends the approved successor SKU with compatibility, price impact and availability",
           "Approve / Modify (type another SKU) / Escalate"],
          ["Line 2: PN-DRAIN-STD (qty 30)",
           "Code is not a valid catalog SKU",
           "Wrong / unrecognised SKU",
           "Identifies the product from its description and proposes the correct SKU (SKU-DRN-3010)",
           "Approve / Type SKU / Reject / Escalate"],
          ["Line 3: (no code) 'Tank-to-Bowl Gasket Kit' (qty 50)",
           "SKU is missing (description + qty only)",
           "Missing SKU",
           "Identifies the product by description and proposes the SKU (SKU-SEL-1150)",
           "Approve / Type SKU / Reject / Escalate"],
          ["Line 4: SKU-VLV-2201 (qty 0)",
           "Quantity is zero / invalid",
           "Invalid quantity",
           "Asks the CSR to enter the correct quantity for the line",
           "Enter qty / Reject / Escalate"],
          ["Line 5: SKU-CTG-4520 (qty 2, no UOM)",
           "No UOM and the qty is small — could be pieces or packs",
           "UOM ambiguity (AC-02)",
           "Shows both readings: 2 EA vs 2 CASE = 48 EA, with the conversion rule",
           "Approve (2 EA default) / Pick packs / Reject / Escalate"],
          ["Ship-to: 'Great Lakes - Ketchikan Project Site' (name only)",
           "Only a partial ship-to (no address / ZIP)",
           "Partial ship-to",
           "Matches it to the registered site (ZIP 99950) and asks for confirmation",
           "Approve / Type address / Reject / Escalate"],
          ["Line 6: SKU-CTG-4520 (qty 2000)",
           "Volume + contract + rebate discount (24.8%) breaches the 18% CARTRIDGE policy",
           "Pricing & Promo exception",
           "'Requested discount 24.8% exceeds policy. Margin impact ≈ $1,700. Approve exception?'",
           "Approve exception / Reject / Escalate (Pricing Desk)"],
          ["Whole order ≈ $38,642 (large order)",
           "Order value exceeds the customer's remaining available credit ($30,000)",
           "Credit hold",
           "'Order value $38,642.12 exceeds available credit $30,000.00. Approve override or send to Finance?'",
           "Approve override / Reject / Escalate (Finance / Credit Team)"],
          ["Line 7: SKU-SHS-7700 (qty 10)",
           "Only 8 on hand — 2 short",
           "Inventory shortage",
           "Proposes a partial shipment now + backorder the rest, with the ETA",
           "Approve partial / Reject / Escalate (Fulfillment Planner)"],
          ["Ship-to resolves to Ketchikan, AK 99950",
           "Remote ZIP not on the standard carrier lanes",
           "Logistics — ZIP not serviceable",
           "Flags the delivery risk and proposes an alternate service",
           "Approve alternate / Reject / Escalate (Logistics Team)"],
      ],
      widths=[1.9, 1.7, 1.3, 2.4, 1.7])

# ── 5. The exact sequence ────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, "5.  Exact CSR gate sequence for CSR-Approval-PO")
body(doc, "When you submit CSR-Approval-PO the AI pauses, in this order, for eleven "
          "CSR decisions — one in every decision layer that supports human-in-the-loop "
          "approval. Click Approve (accept the AI's recommendation) at each one to "
          "walk the full happy-approval path to a completed order.")
table(doc,
      ["#", "Decision layer", "CSR gate", "Recommended demo click"],
      [
          ["1", "Intake", "Unresolved buyer", "Pick 'John Miller'"],
          ["2", "Product Match", "Obsolete-SKU substitution", "Approve (accept substitute)"],
          ["3", "Intake", "Wrong / unrecognised SKU", "Approve (accept SKU-DRN-3010)"],
          ["4", "Intake", "Missing SKU", "Approve (accept SKU-SEL-1150)"],
          ["5", "Intake", "Invalid quantity", "Enter a quantity (e.g. 15)"],
          ["6", "Product Match", "UOM ambiguity", "Approve (2 EA)"],
          ["7", "Customer Validation", "Partial ship-to", "Approve (confirm Ketchikan site)"],
          ["8", "Pricing and Promo", "Discount / margin exception", "Approve exception"],
          ["9", "Credit", "Credit hold (order over available credit)", "Approve override"],
          ["10", "Inventory Checks", "Inventory shortage", "Approve partial + backorder"],
          ["11", "Logistics", "ZIP not serviceable", "Approve alternate"],
      ],
      widths=[0.4, 1.8, 2.3, 2.7])
body(doc, "After the eleventh approval the order completes end-to-end. Final order "
          "total ≈ $38,642.12; the customer confirmation is issued with price, "
          "ETA, fulfillment source and tracking.", )
body(doc, "You can also demonstrate Reject (stops the order) and Escalate (routes "
          "to the responsible team shown in section 4) on any gate instead of "
          "Approve.", italic=True, color=GREY)

# ── 6. Decision-layer gate detail ────────────────────────────────────────────
h1(doc, "6.  Decision-layer gates — the numbers behind them")
h2(doc, "Pricing and Promo (Line 6, SKU-CTG-4520 x 2000)")
bullet(doc, "List price $12.50; negotiated contract price $10.80; volume-tier "
            "discount 10% at 2,000 units plus customer rebates (CARTRIDGE 2% + "
            "loyalty 1%) → net unit ≈ $9.40.")
bullet(doc, "Effective discount vs list ≈ 24.8%, which exceeds the CARTRIDGE policy "
            "limit of 18%.")
bullet(doc, "The AI quantifies the margin impact (≈ $1,700) and asks the CSR to "
            "approve the exception, reject, or escalate to the Pricing Desk.")
h2(doc, "Credit (whole order ≈ $38,642)")
bullet(doc, "Once pricing is approved the AI totals the order at ≈ $38,642.12.")
bullet(doc, "The customer's remaining available credit is $30,000, so the order "
            "value exceeds the available line and the account is placed on credit "
            "hold. (A normal small order — e.g. the happy-flow PO at ≈ $3,349 — "
            "stays well within the line and never triggers this gate.)")
bullet(doc, "The AI asks the CSR to approve a credit override, reject, or escalate "
            "to the Finance / Credit Team.")
h2(doc, "Inventory Checks (Line 7, SKU-SHS-7700 x 10)")
bullet(doc, "8 units on hand at the Chicago DC; the order needs 10 → 2 short.")
bullet(doc, "The AI proposes shipping 8 now and backordering 2 (with the "
            "replenishment ETA), and asks the CSR to approve the partial plan or "
            "escalate to the Fulfillment Planner.")
h2(doc, "Logistics (ship-to Ketchikan, AK 99950)")
bullet(doc, "The confirmed ship-to is a remote Alaska project site whose ZIP is not "
            "on the standard carrier lanes.")
bullet(doc, "The AI flags the delivery risk and asks the CSR to approve an alternate "
            "service or escalate to the Logistics Team.")

# ── 7. Additional scenarios ──────────────────────────────────────────────────
doc.add_page_break()
h1(doc, "7.  Duplicate PO — the one case with no CSR approval")
body(doc, "Every CSR approval gate is now covered by the single CSR-Approval-PO run "
          "(sections 4–6). The only remaining scenario is Duplicate PO, and it is "
          "intentionally NOT a CSR approval — there is nothing for a CSR to approve, "
          "because a duplicate order must never be reprocessed.")
body(doc, "How to demo it: submit Happy-Flow-PO (or CSR-Approval-PO) a second time in "
          "the same session without clearing it. The AI recognises the PO number was "
          "already processed and auto-stops — reprocessing would double-book "
          "inventory and double-charge the customer. No Approve/Override option is "
          "offered; the only action is Escalate (to the Order Operations team) so "
          "they can reconcile against the original submission.")

# ── 8. One PO, every gate ────────────────────────────────────────────────────
h1(doc, "8.  One PO now carries every CSR approval gate")
body(doc, "Earlier, a Credit hold appeared to need a separate credit-risk customer. "
          "That is now solved in master data: the demo customer's available credit is "
          "set below a large-order value, so the big CSR-Approval-PO (≈ $38,642) trips "
          "the Credit gate while ordinary small orders (e.g. the happy-flow PO at "
          "≈ $3,349) still clear credit automatically. As a result the single "
          "CSR-Approval-PO exercises all eleven CSR approval gates in one run.")
bullet(doc, "The CSR file's ship-to does double duty: it is used both for the Partial "
            "ship-to gate and (once confirmed) the Logistics 'ZIP not serviceable' gate.")
bullet(doc, "Duplicate detection needs the same PO submitted twice, so it is a "
            "re-submission demo (section 7) — and it needs no CSR approval anyway.")

# ── 9. Summary ───────────────────────────────────────────────────────────────
h1(doc, "9.  CSR approval coverage summary")
table(doc,
      ["Decision layer", "CSR gate", "Where demoed"],
      [
          ["Intake", "Unresolved buyer", "CSR-Approval-PO (gate 1)"],
          ["Product Match", "Obsolete-SKU substitution", "CSR-Approval-PO (gate 2)"],
          ["Intake", "Wrong / unrecognised SKU", "CSR-Approval-PO (gate 3)"],
          ["Intake", "Missing SKU", "CSR-Approval-PO (gate 4)"],
          ["Intake", "Invalid quantity", "CSR-Approval-PO (gate 5)"],
          ["Product Match", "UOM ambiguity (AC-02)", "CSR-Approval-PO (gate 6)"],
          ["Customer Validation", "Partial ship-to", "CSR-Approval-PO (gate 7)"],
          ["Pricing and Promo", "Discount / margin exception", "CSR-Approval-PO (gate 8)"],
          ["Credit", "Credit hold", "CSR-Approval-PO (gate 9)"],
          ["Inventory Checks", "Inventory shortage / partial fulfillment", "CSR-Approval-PO (gate 10)"],
          ["Logistics", "ZIP not serviceable", "CSR-Approval-PO (gate 11)"],
          ["Intake", "Duplicate PO (NO CSR approval — auto-reject / escalate)", "Section 7 (re-submit)"],
      ],
      widths=[1.8, 2.8, 2.4])

doc.save(OUT)
print("Created:", OUT)
