"""Generate a Word (.docx) document that transcribes the architecture diagram
'Autonomous Purchase Order-to-Fulfillment Orchestration with Human in-the-loop
governance' into a structured written form."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

EY_YELLOW = RGBColor(0xFF, 0xE6, 0x00)
EY_DARK = RGBColor(0x2E, 0x2E, 0x38)
HDR_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)


def shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, lines, bold_first=False, size=9, color=None, fill=None):
    if fill:
        shade_cell(cell, fill)
    cell.text = ""
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(ln)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        if bold_first and i == 0:
            run.bold = True


def heading(doc, text, size=14, color=HDR_BLUE, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(it)
        r.font.size = Pt(size)


doc = Document()

# ── Global default font ───────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Title block ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Autonomous Purchase Order-to-Fulfillment Orchestration")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = EY_DARK
sub = doc.add_paragraph()
sr = sub.add_run("with Human-in-the-Loop Governance")
sr.bold = True
sr.font.size = Pt(16)
sr.font.color.rgb = GREY

# Trigger / scenario line
scen = doc.add_paragraph()
scen.paragraph_format.space_before = Pt(6)
sc = scen.add_run('Customer request (trigger):  "Need 250 units of Faucet SV-200 — Deliver to ZIP 75201"')
sc.italic = True
sc.font.size = Pt(11)
sc.font.color.rgb = HDR_BLUE

doc.add_paragraph()

# ── 1. Customer Interaction Layer ────────────────────────────────────────────
heading(doc, "1. Customer Interaction Layer (Intake Channels)")
doc.add_paragraph(
    "Purchase orders can enter the system through any of the following channels; "
    "the AI normalizes each into a single structured order for downstream processing:"
).runs[0].font.size = Pt(10)
bullets(doc, [
    "Email",
    "PDF PO",
    "Excel PO",
    "Scanned PO",
])

# ── 2. Decision Layer ────────────────────────────────────────────────────────
heading(doc, "2. Decision Layer")
doc.add_paragraph(
    "The Decision Layer runs a sequence of autonomous checks. At any step where a "
    "business rule is breached or data is ambiguous, the agent pauses and asks a CSR "
    "to Approve, Modify, or Escalate before continuing (human-in-the-loop governance)."
).runs[0].font.size = Pt(10)

# Each stage: (title, purpose-label, [purpose items], [example lines], [csr/escalation lines])
stages = [
    ("Intake", "Extracts",
     ["SKU", "Qty", "Ship-to ZIP", "Delivery Date"],
     ["SKU → SV-200", "Qty → 250", "ZIP → 75201"],
     []),
    ("Customer Validation", "Identifies",
     ["Customer tier", "Distributor", "Contractor", "Buying History"],
     ["ABC Plumbing", "Sold Distributor", "Net 45 Terms"],
     []),
    ("Product Match", "Validates",
     ["SKU active", "UOM conversion", "Compatibility", "Buying History"],
     ["AI detects SV-200 is obsolete", "Recommends SV-220"],
     ["Approve substitution?", "CSR actions: Approve / Modify / Escalate"]),
    ("Pricing & Promo", "Determines",
     ["Contract price", "Rebates", "Volume discounts", "Promos", "Freight terms"],
     ["List Price: $120", "Contract Price: $105", "Volume Discount: 5%",
      "Final Price: $99.75"],
     ["If Discount > 10% OR Margin below threshold, AI asks CSR:",
      '"Requested discount exceeds policy. Margin impact = $12K. Approve exception?"']),
    ("Approval", "Evaluates",
     ["Margin rules", "Discount limits", "Approval matrix", "Auto approval"],
     ["Order Value < $100K", "Margin > 15%", "Result: Auto Approval"],
     []),
    ("Credit", "Checks",
     ["Credit limit", "Open invoices", "Payment risk", "Fraud signals"],
     ["Credit Limit: $500K", "Available: $410K", "Decision: PASS"],
     ["If Credit Hold, AI asks CSR:",
      '"Customer exceeded credit limit. Approve override or send to Finance?"']),
    ("Inventory Checks", "Checks",
     ["Plant Stock", "DC Inventory", "In-transit stock", "ATP Inventory"],
     ["Dallas DC: 350", "Houston DC: 150", "Available: 500",
      "Decision: Inventory Available"],
     ["If Shortage, AI proposes: 400 now, 100 next week",
      "CSR approves customer proposal"]),
    ("Logistics", "Determines",
     ["ZIP Serviceability", "Carrier coverage", "Delivery SLA", "ETA Prediction"],
     ["ZIP: 75201", "Carrier: FedEx Freight", "ETA: June 29",
      "Decision: Serviceable"],
     ["If ZIP unsupported, AI suggests nearest pickup location",
      "CSR reviews customer communication"]),
    ("Optimization", "Optimizes",
     ["Warehouse choice", "Shipment split", "Freight cost", "Inventory Balance"],
     ["Option A: Dallas DC — Cost $1,200",
      "Option B: Houston DC — Cost $1,700",
      "Option C: Split Shipment — Cost $1,400"],
     []),
]

for i, (name, label, purpose, example, csr) in enumerate(stages, 1):
    heading(doc, f"2.{i}  {name}", size=12, color=EY_DARK, space_before=8, space_after=2)
    # Purpose
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{label}:  ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(", ".join(purpose)).font.size = Pt(10)
    # Example
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Example: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run("; ".join(example)).font.size = Pt(10)
    # CSR / escalation
    if csr:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("Human-in-the-loop: ")
        r.bold = True
        r.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
        r.font.size = Pt(10)
        for j, line in enumerate(csr):
            pp = p if j == 0 else doc.add_paragraph()
            run = pp.add_run(line if j == 0 else "   " + line)
            run.font.size = Pt(10)
            run.italic = True

# ── 3. Execution Orchestration Layer ─────────────────────────────────────────
heading(doc, "3. Execution Orchestration Layer")
doc.add_paragraph(
    "Once the order clears the Decision Layer, execution is orchestrated across the "
    "downstream enterprise systems:"
).runs[0].font.size = Pt(10)

heading(doc, "3.1  Order (Creates)", size=12, color=EY_DARK, space_before=6, space_after=2)
bullets(doc, ["ERP sales order", "OMS request", "WMS pick ticket", "Shipment order"])

heading(doc, "3.2  Communication (Creates)", size=12, color=EY_DARK, space_before=6, space_after=2)
bullets(doc, ["Order accepted", "Price confirmed", "ETA confirmed", "Tracking shared"])

heading(doc, "3.3  Exception Handling (Resolves)", size=12, color=EY_DARK, space_before=6, space_after=2)
bullets(doc, ["Inventory shortage", "Product obsolete", "ZIP not serviceable",
              "Approval escalation"])

# ── Human-in-the-loop review triggers ────────────────────────────────────────
heading(doc, "4. Human-in-the-Loop Review Triggers")
doc.add_paragraph(
    "The following conditions route the order to a CSR for review/decision:"
).runs[0].font.size = Pt(10)
bullets(doc, [
    "Product Substitution / Matching",
    "Pricing Exception",
    "Credit Hold",
    "Inventory Shortage",
    "Delivery Constraint",
    "Strategic Customer Review",
])

# ── Outcome ──────────────────────────────────────────────────────────────────
heading(doc, "5. Outcome")
p = doc.add_paragraph()
r = p.add_run("CUSTOMER RECEIVES CONFIRMED ORDER")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = HDR_BLUE
p = doc.add_paragraph()
p.add_run("Order #SO-98765   |   ETA: June 29   |   Tracking: Available").font.size = Pt(11)

# ── End-to-end flow summary ──────────────────────────────────────────────────
heading(doc, "6. End-to-End Flow (Summary)")
flow = ("Customer Interaction Layer  →  Decision Layer "
        "(Intake → Customer Validation → Product Match → Pricing & Promo → Approval "
        "→ Credit → Inventory Checks → Logistics → Optimization)  →  Execution "
        "Orchestration Layer (Order → Communication → Exception Handling)  →  "
        "Confirmed Order delivered to the customer.")
doc.add_paragraph(flow).runs[0].font.size = Pt(10)

# Footer note
doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run("EY")
fr.bold = True
fr.font.size = Pt(14)
fr.font.color.rgb = EY_DARK

out = "Autonomous_PO_to_Fulfillment_Orchestration.docx"
doc.save(out)
print("Saved:", out)
