"""Generate the client-facing demo briefing document.

Purpose (per manager's guidance): before the live demo, give the client a short
business perspective, then walk through what each agent/process does and how it
scales up in a real production implementation — all kept short and talk-track
friendly.

Run:  python _build_demo_briefing.py
Output:  demo/Order-Orchestration-Demo-Briefing.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── palette ────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x2A, 0x44)
TEAL   = RGBColor(0x12, 0x7D, 0x7D)
GREY   = RGBColor(0x55, 0x55, 0x60)
DGREY  = RGBColor(0x33, 0x33, 0x3A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
FONT   = "Segoe UI"


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def run(p, text, size=10.5, bold=False, italic=False, color=DGREY, font=FONT):
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = font
    return r


def para(doc, space_before=0, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def heading(doc, text, size=14, color=TEAL, space_before=14, space_after=6):
    p = para(doc, space_before=space_before, space_after=space_after)
    run(p, text, size=size, bold=True, color=color)
    return p


def bullet(doc, runs, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        run(p, text, size=size, bold=bold, color=DGREY)
    return p


# ── document ────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)

normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(10.5); normal.font.color.rgb = DGREY

# Title block
p = para(doc, space_after=2)
run(p, "AI-Powered Order Creation & Orchestration", size=20, bold=True, color=NAVY)
p = para(doc, space_after=2)
run(p, "Client Demo Briefing — Business Perspective & Solution Walkthrough",
    size=12, bold=False, color=TEAL)
p = para(doc, space_after=10)
run(p, "Intelligent PO-to-Order fulfillment with human-in-the-loop governance",
    size=10, italic=True, color=GREY)

# ── 0. How to start the demo (opening talk-track) ────────────────────────────
heading(doc, "0. How to Start the Demo (your opening)")
p = para(doc, space_after=4)
run(p, "Cue \u2014 manager hands over: ", bold=True, color=NAVY)
run(p, "\u201cHi Sanjay, handover to you, please start the demo with a small "
       "introduction.\u201d", italic=True, color=GREY)

p = para(doc, space_before=2, space_after=4)
run(p, "Your opening (suggested talk-track):", bold=True, color=TEAL)

p = para(doc)
run(p, "\u201cHello all, good evening. I am Sanjay and I will walk you through our "
       "AI-powered PO fulfillment POC. We have built a set of AI agents that "
       "automate the purchase order fulfillment process end-to-end, with a human "
       "in the loop to take decisions on exception cases.")
p = para(doc)
run(p, "For this POC we have used realistic, industrial-level data and simulated "
       "enterprise systems like ERP, OMS and PIM through mock APIs, which later can "
       "be replaced with real API integrations. I\u2019ll run two short scenarios "
       "\u2014 first, a clean order that flows straight through, and second, an "
       "exception order where the AI pauses for a quick human decision. Let\u2019s "
       "get started.\u201d")

p = para(doc, space_before=6, space_after=4)
run(p, "Then kick off the live demo:", bold=True, color=TEAL)
bullet(doc, [("Open the app \u2014 ", True), ("launch the PO Fulfillment AI screen.", False)])
bullet(doc, [("Show the architecture (optional) \u2014 ", True), ("one line on the agents and the systems they connect to.", False)])
bullet(doc, [("Run the happy-path PO \u2014 ", True), ("submit a clean order and narrate each agent as it validates and passes straight through.", False)])
bullet(doc, [("Run the exception PO \u2014 ", True), ("submit an order that triggers a CSR decision; show the human-in-the-loop step and resume.", False)])
bullet(doc, [("Close on value \u2014 ", True), ("faster cycle time, fewer errors, and full auditability at scale.", False)])

# 10-minute run sheet
p = para(doc, space_before=8, space_after=4)
run(p, "Keep it to ~10 minutes \u2014 run sheet:", bold=True, color=TEAL)

run_rows = [
    ("0:00 \u2013 1:00", "Intro & handover",
     "Greeting, who you are, one-line context (manual PO-to-order pain).",
     "\u201cThanks. Good morning everyone \u2014 I\u2019m Sanjay, and I\u2019ll walk you through our AI-powered PO fulfillment proof of concept.\u201d"),
    ("1:00 \u2013 2:00", "Business perspective",
     "The \u201cwhy\u201d + the value in a couple of sentences; set up the two scenarios.",
     "\u201cToday, turning a PO into an order is manual across ERP, PIM, Commerce, OMS and shipping. Our agents automate it end-to-end \u2014 people step in only for exceptions.\u201d"),
    ("2:00 \u2013 5:00", "Happy-path PO",
     "Submit a clean order; let it flow straight-through, narrating agents briefly.",
     "\u201cHere\u2019s a clean order \u2014 watch each agent validate against its system and pass it straight through in minutes, with no manual effort.\u201d"),
    ("5:00 \u2013 8:00", "Exception PO",
     "Submit an order that triggers a CSR decision; show human-in-the-loop and resume.",
     "\u201cThis order hits an exception \u2014 the AI pauses, explains the issue, the CSR makes one decision, and the pipeline resumes automatically.\u201d"),
    ("8:00 \u2013 10:00", "Scale-up & close",
     "Mock APIs \u2192 live integrations, scale & governance; close on value, invite questions.",
     "\u201cIn production these mock APIs become live integrations \u2014 same agents, at scale, fully governed. Faster, more accurate, and auditable. Happy to take questions.\u201d"),
]
rtable = doc.add_table(rows=1, cols=4)
rtable.alignment = WD_TABLE_ALIGNMENT.CENTER
rtable.style = "Table Grid"
rwidths = (Inches(0.85), Inches(1.35), Inches(2.35), Inches(2.35))
rhdr = rtable.rows[0].cells
for c, label in zip(rhdr, ("Time", "Segment", "What to cover", "What to say (script)")):
    _shade(c, "1F2A44"); _set_cell_margins(c)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    run(c.paragraphs[0], label, size=9.5, bold=True, color=WHITE)
for r, (t, seg, cover, say) in enumerate(run_rows, start=1):
    cells = rtable.add_row().cells
    if r % 2 == 0:
        for c in cells:
            _shade(c, "F1F5F8")
    for c in cells:
        _set_cell_margins(c)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    run(cells[0].paragraphs[0], t, size=9.5, bold=True, color=TEAL)
    run(cells[1].paragraphs[0], seg, size=9.5, bold=True, color=NAVY)
    run(cells[2].paragraphs[0], cover, size=9.5, color=DGREY)
    run(cells[3].paragraphs[0], say, size=9.5, italic=True, color=GREY)
for row in rtable.rows:
    for i, w in enumerate(rwidths):
        row.cells[i].width = w

p = para(doc, space_before=6)
run(p, "Tip: ", bold=True, color=NAVY)
run(p, "keep the intro tight (~1 min) and let the clean order run quickly \u2014 spend "
       "your richest time on the exception / human-in-the-loop moment, which is the "
       "true differentiator. Leave the last ~2 minutes for scale-up and questions.",
    italic=True, color=GREY)

# ── 1. Business perspective ──────────────────────────────────────────────────
heading(doc, "1. Business Perspective (the \u201cwhy\u201d)")
p = para(doc)
run(p, "Today, turning a customer Purchase Order into a confirmed sales order is "
       "largely manual. Customer Service Reps (CSRs) re-key PO data and chase "
       "validations across many systems \u2014 ERP, PIM, Commerce, OMS and shipping "
       "\u2014 before an order can be created. The result is slow cycle times, "
       "avoidable errors, and cost that grows with volume.")
p = para(doc)
run(p, "This accelerator orchestrates a team of specialised AI agents that "
       "automate the full journey \u2014 intake \u2192 validation \u2192 order creation "
       "\u2014 while keeping a human in the loop for exceptions only. Clean orders "
       "flow ", )
run(p, "straight-through", bold=True, color=NAVY)
run(p, "; only orders with a real exception pause for a quick CSR decision.")

heading(doc, "Business value", size=12, space_before=8, space_after=4)
bullet(doc, [("Faster cycle time \u2014 ", True), ("PO to confirmed order in minutes, not hours.", False)])
bullet(doc, [("Fewer errors \u2014 ", True), ("automated checks catch obsolete SKUs, UOM/pricing mismatches, credit and stock issues.", False)])
bullet(doc, [("Lower cost to serve \u2014 ", True), ("CSRs focus only on genuine exceptions, not routine data entry.", False)])
bullet(doc, [("Governed & auditable \u2014 ", True), ("every decision is explained and logged; humans stay in control.", False)])
bullet(doc, [("Scalable \u2014 ", True), ("the same orchestration handles a handful or thousands of POs a day.", False)])

# ── 2. How the demo works ────────────────────────────────────────────────────
heading(doc, "2. How the Solution Works \u2014 What Each Process Does")
p = para(doc)
run(p, "The pipeline is a sequence of specialised agents. Each agent connects to "
       "the relevant enterprise system(s), validates its part of the order, and "
       "hands off to the next. In the demo the systems are simulated by mock APIs "
       "using real master data; in production they become live integrations.")

# Agent table: name | what it does (business) | systems | scale-up
rows = [
    ("Intake Agent",
     "Reads the incoming PO (email / Excel) and extracts header and line items into structured data \u2014 no manual re-keying.",
     "PO document",
     "Auto-ingest from Outlook / Graph, EDI, customer portals and PDF/OCR at high volume."),
    ("Customer Validation Agent",
     "Confirms the customer account, ship-to and that the buyer is authorised to place the order.",
     "Commerce, OMS, PIM",
     "Live Commerce/CRM APIs and an enterprise buyer directory (SSO) with fraud controls."),
    ("Product Matching Agent",
     "Validates SKUs, product lifecycle (e.g. obsolete), unit-of-measure (KIT vs EA) and compliance.",
     "PIM, ERP",
     "Real PIM / product master (MDM) with automated substitution suggestions."),
    ("Pricing & Promo Agent",
     "Applies contract pricing and promotions, and checks margin before the order proceeds.",
     "Commerce",
     "Enterprise pricing / CPQ and contract-management engines."),
    ("Credit Agent",
     "Checks credit limit, payment terms and account risk to protect revenue.",
     "ERP",
     "ERP accounts-receivable plus external credit-risk services."),
    ("Inventory Agent",
     "Confirms stock availability across distribution centres so the order can actually be fulfilled.",
     "ERP",
     "Real-time available-to-promise (ATP) across DCs and WMS."),
    ("Shipments Agent",
     "Determines the carrier, service level and delivery SLA that fit the order.",
     "Shipping Provider, ERP",
     "TMS and live carrier APIs for rating, booking and tracking."),
    ("Optimization Agent",
     "Selects the best fulfillment plan \u2014 balancing cost, speed and shipment splits.",
     "Shipping Provider, ERP",
     "Optimisation engine spanning multiple DCs, carriers and cost constraints."),
    ("Approvals Agent",
     "Runs the approvals matrix and budget checks so the order meets policy.",
     "ERP",
     "Enterprise workflow / approval systems with delegation rules."),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
widths = (Inches(1.55), Inches(3.05), Inches(1.15), Inches(2.6))
hdr = table.rows[0].cells
for c, label in zip(hdr, ("Agent / Process", "What it does (business view)", "Systems", "Scale-up in real implementation")):
    _shade(c, "1F2A44"); _set_cell_margins(c)
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    run(c.paragraphs[0], label, size=9.5, bold=True, color=WHITE)
for r, (name, does, sysx, scale) in enumerate(rows, start=1):
    cells = table.add_row().cells
    if r % 2 == 0:
        for c in cells:
            _shade(c, "F1F5F8")
    for c in cells:
        _set_cell_margins(c)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    run(cells[0].paragraphs[0], name, size=9.5, bold=True, color=NAVY)
    run(cells[1].paragraphs[0], does, size=9.5, color=DGREY)
    run(cells[2].paragraphs[0], sysx, size=9.5, bold=True, color=TEAL)
    run(cells[3].paragraphs[0], scale, size=9.5, color=GREY)
for row in table.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w

# ── 3. Human-in-the-loop ─────────────────────────────────────────────────────
heading(doc, "3. Human-in-the-Loop Governance")
p = para(doc)
run(p, "Every agent can pause on an exception and hand the decision to a CSR "
       "(approve, reject, correct, escalate or choose an option). If there is no "
       "exception, the order passes straight through. This keeps automation fast "
       "while people stay firmly in control of judgement calls \u2014 and every "
       "action is captured in an audit trail.")

# ── 4. Order execution ───────────────────────────────────────────────────────
heading(doc, "4. Order Creation & Downstream")
p = para(doc)
run(p, "Once all checks pass, the Order Execution step creates the sales order and "
       "triggers the downstream flow:")
bullet(doc, [("ERP \u2014 ", True), ("sales order created", False)])
bullet(doc, [("OMS \u2014 ", True), ("fulfillment request", False)])
bullet(doc, [("WMS \u2014 ", True), ("pick ticket", False)])
bullet(doc, [("TMS \u2014 ", True), ("shipment booked & tracked", False)])
bullet(doc, [("Notification \u2014 ", True), ("confirmation email to the customer", False)])

# ── 5. Scale-up ──────────────────────────────────────────────────────────────
heading(doc, "5. From Demo to Production \u2014 How It Scales")
bullet(doc, [("Same orchestration, real integrations \u2014 ", True),
             ("the demo\u2019s mock APIs are swapped for live ERP / PIM / Commerce / OMS / TMS APIs; the agents and their logic stay the same.", False)])
bullet(doc, [("Multi-channel intake \u2014 ", True),
             ("email, EDI, portals and PDFs ingested automatically at scale.", False)])
bullet(doc, [("Event-driven & high-volume \u2014 ", True),
             ("processes thousands of POs per day with parallel agent execution.", False)])
bullet(doc, [("Enterprise-grade \u2014 ", True),
             ("security, role-based access, monitoring, and a full audit trail for compliance.", False)])
bullet(doc, [("Continuous improvement \u2014 ", True),
             ("exception patterns feed back to reduce future manual touches over time.", False)])

p = para(doc, space_before=12)
run(p, "In short: ", bold=True, color=NAVY)
run(p, "the demo shows the end-to-end experience on realistic data; the production "
       "path simply connects the same agents to your live systems \u2014 delivering "
       "faster, more accurate, and fully governed order fulfillment at scale.",
    color=DGREY)

out = os.path.join("demo", "Order-Orchestration-Demo-Briefing.docx")
doc.save(out)
print("Saved:", out)
