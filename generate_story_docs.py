"""
generate_story_docs.py
Generates a manager summary Word document for each user story US-03 ... US-12,
placed inside its own folder: sample-data/US-XX/US-XX_Summary_for_Manager.docx

Each document mirrors US-02_Summary_for_Manager.docx and contains:
  - Story overview (layer, capability, persona, business value, responsibilities,
    systems/data needed, exception types)
  - Every acceptance criterion (full Given/When/Then) and how the POC satisfies it
  - What was built (module, master-data workbook + sheets/fields, how it works)
  - Exception scenarios mapped to the sample PO files in the same folder
  - How to demo

Run:  python generate_story_docs.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.join(os.path.dirname(__file__), "sample-data")
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
BLUE = RGBColor(0x25, 0x63, 0xEB)

# ── Story content ─────────────────────────────────────────────────────────────
STORIES = {
    "US-03": {
        "capability": "Buyer Authorization, Product Visibility, and Ordering Rights",
        "layer": "Decision Layer / Customer Validation",
        "persona": "B2B Buyer / CSR / Account Manager",
        "user_story": ("As a B2B platform user, I want the AI agent to validate whether the buyer is "
                       "authorized to purchase requested products for the specified branch, cost center, "
                       "or ship-to location, so that orders comply with account-level purchasing policies."),
        "business_value": ("Prevents unauthorized purchases, supports role-based procurement controls, "
                           "and enforces customer-specific catalog visibility."),
        "ai_resp": "Validate buyer role, branch, cost center, product visibility, order limits, and purchasing rights.",
        "human_resp": "Review authorization exceptions and coordinate with account managers or customer administrators.",
        "systems": "Buyer profile, user permissions, account hierarchy, cost center rules, product visibility rules, punchout/eProcurement mappings.",
        "exceptions": ["Unauthorized buyer", "Restricted product", "Invalid cost center", "Order threshold violation"],
        "module": "modules/buyer_authorization.py",
        "master": "buyer-master-data.xlsx",
        "sheets": [
            ("Buyer_Profiles", "buyer_id, buyer_name, email, customer_account, branch_id, default_cost_center, role, status, max_order_value, currency, can_self_approve, punchout_id"),
            ("User_Permissions", "buyer_id, permitted_branches, permitted_cost_centers, allowed_product_families, denied_product_families, max_line_value, requires_approval_above"),
            ("Cost_Centers", "cost_center_id, name, branch_id, status, owner, currency"),
            ("Product_Visibility_Rules", "rule_id, scope_type, scope_id, product_family, sku, visibility, min_role, reason"),
        ],
        "how": ("The agent resolves the buyer from Buyer_Profiles, checks the buyer is ACTIVE and tied to the "
                "order's customer/branch, validates the cost center (exists, ACTIVE, owned by the order branch, and "
                "within the buyer's permitted scope), then evaluates product visibility for each requested SKU using "
                "the buyer's denied families and the Product_Visibility_Rules (role-gated and cost-center-hidden items)."),
        "acs": [
            ("AC-01", "Validate buyer ordering rights",
             "Given the customer account hierarchy has been identified\nWhen the AI agent checks the buyer profile, branch, cost center, and assigned permissions\nThen the system should confirm whether the buyer is authorized to submit the order\nAnd validate that the requested products are visible and orderable for the buyer's assigned account level\nAnd proceed to product validation if authorization is successful",
             "The buyer is matched and status/branch/cost-center scope are confirmed; visible products proceed to product validation."),
            ("AC-02", "Restrict product not visible to branch or cost center",
             "Given a buyer submits a purchase order for a restricted product\nWhen the AI agent validates product visibility rules\nThen the system should create a product visibility exception\nAnd display the restricted SKU, applicable visibility rule, and impacted hierarchy level\nAnd route the exception to the CSR or account manager for review",
             "A RESTRICTED_PRODUCT exception lists each blocked SKU with the reason and rule level (sample: scenario-restricted-product.txt)."),
            ("AC-03", "Validate cost center assignment",
             "Given a purchase order references a cost center\nWhen the AI agent validates the cost center against the buyer and account hierarchy\nThen the system should confirm that the cost center is active and belongs to the buyer's permitted branch or division\nAnd create an authorization exception if the cost center is invalid, inactive, or outside the buyer's permission scope",
             "An INVALID_COST_CENTER exception fires when the cost center is missing, inactive, or out of branch/scope (sample: scenario-invalid-cost-center.txt)."),
        ],
        "scenarios": [
            ("Unauthorized buyer", "scenario-unauthorized-buyer.txt", "UNAUTHORIZED_BUYER (suspended buyer BUY-900)"),
            ("Restricted product", "scenario-restricted-product.txt", "RESTRICTED_PRODUCT (junior buyer ordering a denied valve)"),
            ("Invalid cost center", "scenario-invalid-cost-center.txt", "INVALID_COST_CENTER (inactive CC-OLD-900)"),
            ("Happy path", "happy-path.txt", "Authorized — proceeds to product validation"),
        ],
    },
    "US-04": {
        "capability": "Complex Product Matching, Configuration, Variant, and UOM Validation",
        "layer": "Decision Layer / Product Match",
        "persona": "CSR / Product Specialist / Order Operations User",
        "user_story": ("As a CSR or order operations user, I want the AI agent to validate complex B2B product "
                       "attributes, configurations, variants, and units of measure, so that the correct industrial "
                       "product variant is matched before pricing and fulfillment."),
        "business_value": ("Improves product match accuracy for complex catalogs and reduces fulfillment errors "
                           "caused by incorrect variants or UOM conversion issues."),
        "ai_resp": "Match SKU/description to catalog variant, validate required attributes, convert UOM, and recommend substitutes for obsolete/inactive products.",
        "human_resp": "Confirm ambiguous matches, incomplete configurations, substitutions, or non-standard UOM requests.",
        "systems": "Product master, configurable product attributes, compatibility rules, UOM conversion table, substitution rules, buying history.",
        "exceptions": ["Obsolete SKU", "Inactive SKU", "Ambiguous product description", "Missing attributes", "Invalid UOM", "Substitution approval"],
        "module": "modules/product_matcher.py",
        "master": "product-master-data.xlsx",
        "sheets": [
            ("Product_Master", "sku, description, product_family, status, base_uom, material, grade, size, configurable, hazardous, substitute_sku, list_price, currency, weight_kg, lead_time_days, manufacturer, country_of_origin"),
            ("Product_Attributes", "sku, attribute_name, attribute_value, required, uom_dimension"),
            ("UOM_Conversions", "from_uom, to_uom, factor, product_family, notes"),
            ("Substitution_Rules", "original_sku, substitute_sku, compatibility, price_impact_pct, availability_impact, requires_approval, rationale"),
            ("Compatibility_Rules", "sku, compatible_with_sku, rule_type, notes"),
        ],
        "how": ("Each line's SKU is matched to a catalog variant; unmatched SKUs raise a configuration exception with "
                "possible matches. OBSOLETE/INACTIVE SKUs produce an approved-substitute recommendation. Non-standard "
                "units are converted to the base UOM using approved conversion factors, and an invalid-UOM exception is "
                "raised when no conversion rule exists."),
        "acs": [
            ("AC-01", "Match requested SKU to configured product variant",
             "Given a purchase order contains a product with multiple configurable attributes\nWhen the AI agent validates the product against the catalog\nThen the system should match the requested SKU or description to the correct product variant\nAnd validate required attributes such as size, material, grade, compatibility, region, and configuration options\nAnd proceed only when the product configuration is complete and orderable",
             "SKUs are matched against Product_Master with attributes from Product_Attributes; complete, orderable variants proceed."),
            ("AC-02", "Convert and validate unit of measure",
             "Given a purchase order contains quantity in a non-standard unit of measure\nWhen the AI agent validates the requested product quantity\nThen the system should convert the requested unit of measure using approved conversion rules\nAnd display the original quantity, converted quantity, converted UOM, and conversion logic\nAnd proceed to pricing if the converted quantity is valid",
             "UOM_Conversions are applied and the conversion logic is shown; an INVALID_UOM exception fires when no rule exists (sample: scenario-invalid-uom.txt)."),
            ("AC-03", "Create exception for incomplete or ambiguous product configuration",
             "Given a purchase order contains an incomplete product description or missing configuration attributes\nWhen the AI agent cannot confidently match the product to an orderable variant\nThen the system should create a product configuration exception\nAnd display the missing attributes or possible product matches\nAnd allow the CSR to confirm, modify, or escalate the product match",
             "A PRODUCT_CONFIG_EXCEPTION lists possible matches for unmatched SKUs (sample: scenario-unknown-sku.txt)."),
            ("AC-04", "Recommend substitute for obsolete or inactive SKU",
             "Given the purchase order contains an obsolete or inactive SKU\nWhen the AI agent identifies an approved substitute product\nThen the system should create a substitution recommendation\nAnd display the original SKU, recommended substitute SKU, compatibility rationale, price impact, and availability impact\nAnd require CSR approval before replacing the requested SKU",
             "An OBSOLETE_SKU recommendation shows substitute, compatibility, price/availability impact, and requires approval (sample: scenario-obsolete-sku.txt)."),
        ],
        "scenarios": [
            ("Obsolete SKU", "scenario-obsolete-sku.txt", "OBSOLETE_SKU + substitution recommendation"),
            ("Invalid UOM", "scenario-invalid-uom.txt", "INVALID_UOM (KG cannot convert to FT)"),
            ("Ambiguous / unknown SKU", "scenario-unknown-sku.txt", "PRODUCT_CONFIG_EXCEPTION with possible matches"),
        ],
    },
    "US-05": {
        "capability": "Regional Compliance, Product Eligibility, and SDS Validation",
        "layer": "Decision Layer / Product Match and Compliance",
        "persona": "Compliance User / CSR / Order Operations User",
        "user_story": ("As a compliance or order operations user, I want the AI agent to validate regional product "
                       "eligibility, compliance restrictions, and safety documentation requirements, so that restricted "
                       "or regulated products are not fulfilled incorrectly."),
        "business_value": "Reduces compliance risk and ensures regulated products ship only to eligible locations with required documentation.",
        "ai_resp": "Check product-region eligibility, identify required SDS/compliance documents, attach documentation, and block restricted combinations.",
        "human_resp": "Review compliance exceptions, approve compliant alternatives, or reject restricted orders.",
        "systems": "Compliance rules, regional restrictions, SDS repository, product eligibility master, customer communication templates.",
        "exceptions": ["Restricted product-region combination", "Missing SDS", "Compliance approval required"],
        "module": "modules/compliance_validator.py",
        "master": "compliance-master-data.xlsx",
        "sheets": [
            ("Compliance_Rules", "rule_id, sku, product_family, restriction_type, region_state, status, reason, approval_required, authority"),
            ("Regional_Restrictions", "region_code, region_name, restricted_families, notes"),
            ("SDS_Repository", "sku, sds_document_id, sds_version, issue_date, expiry_date, hazard_class, required, file_ref"),
            ("Product_Eligibility", "sku, region, eligible, conditions"),
        ],
        "how": ("The ship-to ZIP resolves to a region. Each product's eligibility is checked against Product_Eligibility "
                "and Regional_Restrictions; ineligible combinations are blocked. Hazardous products must have an SDS in the "
                "SDS_Repository, which is attached to the order and logged; a missing SDS raises an exception."),
        "acs": [
            ("AC-01", "Validate regional compliance",
             "Given a purchase order includes products subject to regional restrictions\nWhen the AI agent validates the product against the ship-to location\nThen the system should determine whether each product is eligible for sale and shipment to that region\nAnd validate whether required compliance documents are available\nAnd proceed only when compliance validation is successful",
             "Eligibility is resolved per region; eligible lines proceed, restricted ones are blocked."),
            ("AC-02", "Attach SDS or compliance documentation",
             "Given a product requires safety documentation or regulatory attachments\nWhen the AI agent validates the product and fulfillment region\nThen the system should identify the required SDS or compliance document\nAnd attach the document to the order record or customer communication\nAnd log the compliance validation result in the audit trail",
             "Hazardous SKUs have their SDS identified, attached, and logged; carried into the order confirmation."),
            ("AC-03", "Block restricted product-region combination",
             "Given a product is restricted in the ship-to region\nWhen the AI agent completes compliance validation\nThen the system should create a compliance exception\nAnd prevent the order line from progressing to fulfillment\nAnd route the exception to the CSR or compliance approver",
             "A COMPLIANCE_RESTRICTION exception blocks the line and routes to the compliance approver (sample: scenario-restricted-region.txt)."),
        ],
        "scenarios": [
            ("Restricted product-region", "scenario-restricted-region.txt", "COMPLIANCE_RESTRICTION (chemical to CA)"),
            ("Missing SDS", "scenario-missing-sds.txt", "MISSING_SDS (hazardous SKU with no SDS on file)"),
        ],
    },
    "US-06": {
        "capability": "Enterprise B2B Pricing Engine Validation",
        "layer": "Decision Layer / Pricing and Promo",
        "persona": "Pricing User / CSR / Order Operations User",
        "user_story": ("As a pricing or order operations user, I want the AI agent to calculate the final B2B price using "
                       "customer-specific, contract-based, date-bound, volume-based, location-based, and dynamic pricing "
                       "rules, so that enterprise orders are priced accurately and transparently."),
        "business_value": "Supports enterprise-grade B2B pricing complexity and prevents margin leakage from incorrect discounts or expired contract rates.",
        "ai_resp": "Evaluate price list, contract price, validity dates, volume discounts, rebates, surcharges, freight terms, raw-material adjustments, margin thresholds, and pricing exceptions.",
        "human_resp": "Review pricing exceptions, approve margin exceptions, and correct invalid contract references.",
        "systems": "Pricing engine, contract repository, rebate rules, promotion rules, surcharge tables, freight terms, margin policy, raw material index.",
        "exceptions": ["Expired contract", "Invalid contract", "Excess discount", "Margin below threshold", "Missing pricing rule"],
        "module": "modules/pricing_engine.py",
        "master": "pricing-master-data.xlsx",
        "sheets": [
            ("Price_List", "sku, currency, list_price, uom, effective_from, effective_to, price_list_id"),
            ("Contracts", "contract_reference, customer_account, scope_type, scope_id, contract_price, currency, valid_from, valid_to, status"),
            ("Volume_Tiers", "tier_id, scope_type, scope_id, min_qty, max_qty, discount_pct, uom"),
            ("Rebates", "rebate_id, customer_account, product_family, rebate_pct, condition"),
            ("Promotions", "promo_id, scope_type, scope_id, discount_pct, valid_from, valid_to, description"),
            ("Surcharges", "surcharge_id, scope_type, scope_id, surcharge_type, amount_type, amount, reason"),
            ("Freight_Terms", "scope_type, scope_id, incoterm, freight_payer, base_freight, notes"),
            ("Margin_Policy", "product_family, min_margin_pct, max_discount_pct, approver_role"),
            ("Raw_Material_Index", "material, index_date, index_pct_adjustment, notes"),
        ],
        "how": ("For each line the engine selects the base price (date-bound contract price if active, else list price), "
                "applies volume-tier, promotion, and rebate discounts, then adds order-level surcharges and freight. "
                "If the effective discount versus list price breaches the family margin/discount policy, a pricing "
                "exception is raised and routed to the pricing approver."),
        "acs": [
            ("AC-01", "Calculate final price using multi-layer B2B pricing rules",
             "Given the customer account, hierarchy level, product, quantity, UOM, ship-to location, and requested delivery date are validated\nWhen the AI agent invokes the pricing engine\nThen the system should evaluate price list, contract price, validity dates, volume discounts, rebates, promotions, surcharges, freight terms, and raw-material adjustments\nAnd calculate the final price for each order line\nAnd display a pricing breakdown with applied rules and price sources",
             "A per-line breakdown shows list, base source, discounts, net unit, and line total, plus order subtotal/surcharges/freight."),
            ("AC-02", "Apply date-bound contract pricing",
             "Given a customer has a negotiated contract price for a product\nWhen the AI agent evaluates pricing for the requested date\nThen the system should apply the contract price only if the contract is active for the date range\nAnd fall back to the next applicable pricing rule if the contract is expired or not yet active\nAnd log the pricing decision in the order audit trail",
             "Contracts are validated by status and date; expired/inactive contracts fall back to list price with an audit note."),
            ("AC-03", "Apply volume-based tier discounts",
             "Given the requested quantity qualifies for a volume pricing tier\nWhen the AI agent calculates pricing\nThen the system should apply the correct tiered discount based on converted quantity and eligible product family or SKU\nAnd show the tier threshold, applied discount, and final unit price",
             "Volume_Tiers are matched on the converted quantity and family/SKU and reflected in the breakdown."),
            ("AC-04", "Apply location-specific surcharge",
             "Given the ship-to location has a valid location-based surcharge or regional fee\nWhen the AI agent calculates the final price\nThen the system should include the surcharge in the order pricing\nAnd display the surcharge type, amount, and reason",
             "Surcharges (ZIP-specific and global, e.g. fuel) are added and itemized in a surcharges table."),
            ("AC-05", "Trigger pricing exception for margin or discount breach",
             "Given the calculated price includes discounts, rebates, or adjustments\nWhen the AI agent determines the discount exceeds policy or margin falls below threshold\nThen the system should create a pricing exception\nAnd display list price, contract price, applied discount, margin impact, and recommended action\nAnd route the exception to the appropriate pricing approver",
             "A PRICING_EXCEPTION fires when the effective discount exceeds the family policy limit (sample: scenario-pricing-exception.txt)."),
        ],
        "scenarios": [
            ("Excess discount / margin breach", "scenario-pricing-exception.txt", "PRICING_EXCEPTION (high-volume contract discount > policy)"),
            ("Happy path", "happy-path.txt", "Priced within policy; order total calculated"),
        ],
    },
    "US-07": {
        "capability": "Budget, Spend Limit, and Approval Routing Validation",
        "layer": "Decision Layer / Approval",
        "persona": "Buyer Manager / Approver / CSR",
        "user_story": ("As a B2B buyer manager or approver, I want the AI agent to validate order value against budgets "
                       "and spend thresholds across the customer hierarchy, so that orders are routed for approval based "
                       "on enterprise procurement policies."),
        "business_value": "Enforces spend management across global parent, regional division, local branch, cost center, and ship-to levels.",
        "ai_resp": "Calculate order value, check budgets, evaluate buyer approval authority, identify approver, and create approval task.",
        "human_resp": "Approve, reject, or request changes for orders that exceed budgets or approval limits.",
        "systems": "Budget master, cost centers, approval matrix, customer hierarchy, buyer authority rules, workflow engine.",
        "exceptions": ["Budget exceeded", "Spend threshold exceeded", "Missing approver", "Approval timeout"],
        "module": "modules/budget_approval.py",
        "master": "budget-master-data.xlsx",
        "sheets": [
            ("Budget_Master", "level_type, level_id, period, budget_amount, consumed_amount, available_amount, currency"),
            ("Cost_Centers", "cost_center_id, name, branch_id, status, budget_amount, consumed_amount, available_amount, currency"),
            ("Approval_Matrix", "level_type, level_id, min_amount, max_amount, approver_role, approver_name, sla_hours"),
            ("Buyer_Authority", "buyer_id, max_order_value, can_self_approve, branch_id, cost_centers"),
        ],
        "how": ("The order value is compared to the cost-center and branch available budget. If within budget and the "
                "buyer's self-approval authority, the order is auto-approved. Otherwise the approval matrix identifies "
                "the correct approver (branch -> regional -> global) and an approval task is created, pausing the order. "
                "When approval is required from an approver other than the CSR, the system shows a mocked notification - "
                "'Triggered email to respective approver and awaiting approval.' - and then halts: no further stages run "
                "until the approval is granted or rejected."),
        "acs": [
            ("AC-01", "Validate order against branch-level budget",
             "Given an order is associated with a local branch or cost center\nWhen the AI agent calculates the total order value\nThen the system should compare the order value against the available branch-level or cost-center budget\nAnd determine whether the buyer is authorized to submit without approval\nAnd proceed to credit validation if within budget and approval limits",
             "Cost-center then branch budgets are checked; within-authority orders auto-approve and proceed to credit."),
            ("AC-02", "Route order to regional approver when spend threshold is exceeded",
             "Given a local branch user submits an order above the allowed spend threshold\nWhen the AI agent evaluates the approval matrix\nThen the system should identify the correct regional or corporate approver\nAnd create an approval task with order value, budget impact, products, pricing summary, and recommended action\nAnd pause order progression until approval is granted",
             "An APPROVAL_REQUIRED task names the approver, level, threshold range, and SLA. Because approval is needed "
             "from an approver other than the CSR, the system displays a mocked message 'Triggered email to respective "
             "approver and awaiting approval.' and then stops - no further actions execute until a response is received "
             "(sample: scenario-approval-required.txt)."),
            ("AC-03", "Block or escalate order when budget is unavailable",
             "Given the branch or cost center does not have sufficient remaining budget\nWhen the AI agent performs budget validation\nThen the system should create a budget exception\nAnd display available budget, order value, budget shortfall, and impacted account level\nAnd route the exception according to the customer's approval policy",
             "A BUDGET_EXCEEDED exception shows the shortfall and impacted level (sample: scenario-budget-exceeded.txt)."),
            ("AC-04", "Record approval outcome",
             "Given an approval task has been routed to an approver\nWhen the approver approves, rejects, or requests changes\nThen the system should record the approver, timestamp, decision, comments, and approval level\nAnd resume or stop orchestration based on the approval outcome",
             "The approval task captures approver, level, and SLA; on approval the order resumes (human-in-the-loop)."),
        ],
        "scenarios": [
            ("Budget exceeded", "scenario-budget-exceeded.txt", "BUDGET_EXCEEDED (order > cost-center & branch budget)"),
            ("Approval required", "scenario-approval-required.txt", "APPROVAL_REQUIRED -> mocked approver email sent, process halted"),
        ],
    },
    "US-08": {
        "capability": "Credit, Payment Terms, and Financial Risk Validation",
        "layer": "Decision Layer / Credit",
        "persona": "Finance User / CSR / Order Operations User",
        "user_story": ("As a finance or CSR user, I want the AI agent to validate customer credit limit, payment terms, "
                       "open invoices, and financial risk indicators, so that the order is released only when the customer "
                       "is financially eligible."),
        "business_value": "Protects revenue and reduces financial exposure by preventing release when credit or payment risk is unacceptable.",
        "ai_resp": "Check credit limit, available credit, open invoices, payment terms, overdue invoices, fraud signals, and risk indicators.",
        "human_resp": "Review credit hold, approve override, request payment, or route to finance.",
        "systems": "ERP AR, credit master, invoice aging, payment history, finance approval workflow, fraud/risk signals.",
        "exceptions": ["Credit limit exceeded", "Overdue invoices", "Payment risk", "Finance approval required"],
        "module": "modules/credit_validator.py",
        "master": "credit-master-data.xlsx",
        "sheets": [
            ("Credit_Master", "customer_account, credit_limit, available_credit, currency, payment_terms, credit_status, risk_rating"),
            ("Invoice_Aging", "customer_account, invoice_no, amount, due_date, days_overdue, status"),
            ("Payment_History", "customer_account, avg_days_to_pay, last_payment_date, last_payment_amount, dispute_count"),
            ("Payment_Terms", "terms_code, description, net_days, prepayment_required"),
            ("Risk_Signals", "customer_account, fraud_flag, risk_score, watchlist, notes"),
        ],
        "how": ("The agent checks credit status, available credit versus order value, overdue invoices, and fraud/watchlist "
                "signals. Within-policy customers pass with the correct payment terms applied; otherwise the order is placed "
                "on credit hold with the recommended finance action."),
        "acs": [
            ("AC-01", "Approve credit check for eligible customer",
             "Given the order has passed pricing and approval validation\nWhen the AI agent checks credit limit, open invoices, available credit, payment terms, and payment risk\nThen the system should mark the credit check as passed if within policy\nAnd proceed to inventory availability validation",
             "Within-policy customers pass and proceed to inventory; payment terms are applied."),
            ("AC-02", "Place order on credit hold",
             "Given the order value exceeds available credit or the customer has overdue invoices\nWhen the AI agent performs credit validation\nThen the system should place the order on credit hold\nAnd create a credit exception\nAnd display credit limit, available credit, overdue invoice amount, payment terms, and recommended next action\nAnd route the exception to finance or CSR review",
             "A CREDIT_HOLD exception details limit, available credit, overdue invoices, and action (sample: scenario-credit-hold.txt)."),
            ("AC-03", "Respect payment terms",
             "Given a customer account has specific payment terms such as Net 30, Net 45, prepayment, or credit-hold rules\nWhen the AI agent validates the order\nThen the system should apply the correct payment terms from the customer hierarchy or contract\nAnd include the applied payment terms in the order confirmation and audit trail",
             "Payment terms from Credit_Master/Payment_Terms are applied and flow into the order confirmation."),
        ],
        "scenarios": [
            ("Credit hold", "scenario-credit-hold.txt", "CREDIT_HOLD (overdue invoices / watch status)"),
        ],
    },
    "US-09": {
        "capability": "Inventory Availability, ATP, Allocation, and Partial Fulfillment Planning",
        "layer": "Decision Layer / Inventory Checks",
        "persona": "Fulfillment Planner / CSR / Order Operations User",
        "user_story": ("As a fulfillment planner, I want the AI agent to validate inventory availability across plants, "
                       "distribution centers, in-transit inventory, and available-to-promise supply, so that the system "
                       "can determine whether the requested B2B quantity can be fulfilled."),
        "business_value": "Improves promise accuracy, supports partial fulfillment decisions, and balances inventory across the network.",
        "ai_resp": "Check plant/DC/in-transit stock, ATP, allocation rules, backorder tolerance, and customer fulfillment preferences.",
        "human_resp": "Approve partial fulfillment proposals, customer communication, backorder decisions, or allocation exceptions.",
        "systems": "ERP inventory, WMS, ATP engine, allocation rules, customer fulfillment preferences, in-transit inventory feed.",
        "exceptions": ["Inventory shortage", "Allocation conflict", "Backorder approval", "Split fulfillment approval"],
        "module": "modules/inventory_validator.py",
        "master": "inventory-master-data.xlsx",
        "sheets": [
            ("Plant_Stock", "location_id, location_type, region, sku, on_hand_qty, uom"),
            ("DC_Stock", "location_id, location_type, region, sku, on_hand_qty, uom"),
            ("In_Transit", "sku, from_location, to_location, qty, uom, eta_date"),
            ("ATP", "sku, atp_qty, uom, next_replenishment_date, replenishment_qty"),
            ("Allocation_Rules", "customer_tier, priority, backorder_tolerance_days, notes"),
            ("Fulfillment_Preferences", "customer_account, customer_tier, preferred_warehouse, restricted_dc, split_shipment, backorder_tolerance_days"),
        ],
        "how": ("The agent compares requested quantity to available-to-promise and DC stock (honoring preferred and "
                "restricted warehouses), selecting a source. When the full quantity is unavailable it proposes a partial "
                "fulfillment plan with backorder quantity and estimated availability, applying tier-based allocation rules."),
        "acs": [
            ("AC-01", "Confirm inventory availability across fulfillment network",
             "Given an order has passed customer, product, pricing, approval, and credit validation\nWhen the AI agent checks inventory across plant stock, DC inventory, in-transit stock, and ATP\nThen the system should determine whether the requested quantity is available\nAnd identify the source location(s) that can fulfill the order\nAnd proceed to logistics validation",
             "Availability is confirmed against ATP/DC stock and a source is identified per line, then logistics runs."),
            ("AC-02", "Propose allocation when requested quantity is partially available",
             "Given the full requested quantity is not available\nWhen the AI agent identifies partial availability\nThen the system should propose available quantity, backordered quantity, estimated availability date, and fulfillment source\nAnd create an inventory shortage exception\nAnd route the proposal to the CSR for customer confirmation",
             "An INVENTORY_SHORTAGE proposal shows available, backordered, and ETA (sample: scenario-inventory-shortage.txt)."),
            ("AC-03", "Respect customer-specific fulfillment rules",
             "Given the customer account has fulfillment preferences or restrictions\nWhen the AI agent evaluates inventory options\nThen the system should apply rules such as preferred warehouse, restricted DC, split-shipment preference, and backorder tolerance\nAnd include the applied fulfillment rules in the decision rationale",
             "Preferred/restricted warehouses, split-shipment, and backorder tolerance from Fulfillment_Preferences are applied and shown."),
            ("AC-04", "Handle allocation priority",
             "Given inventory is constrained and multiple customers may require the same inventory\nWhen the AI agent evaluates allocation rules\nThen the system should apply allocation priority based on customer tier, contract commitment, order urgency, and configured business rules\nAnd create an allocation exception if the requested quantity cannot be allocated",
             "Allocation_Rules assign priority by customer tier; constrained allocations surface as exceptions."),
        ],
        "scenarios": [
            ("Inventory shortage", "scenario-inventory-shortage.txt", "INVENTORY_SHORTAGE + partial fulfillment proposal"),
        ],
    },
    "US-10": {
        "capability": "Logistics, ZIP Serviceability, Delivery SLA, and Fulfillment Optimization",
        "layer": "Decision Layer / Logistics and Optimization",
        "persona": "Logistics User / CSR / Fulfillment Planner",
        "user_story": ("As a logistics or order operations user, I want the AI agent to validate carrier serviceability, "
                       "delivery SLA, freight cost, warehouse selection, shipment split, and ETA, so that the order is "
                       "fulfilled using the best cost-service option."),
        "business_value": "Optimizes service level, freight cost, inventory balance, and customer delivery commitments.",
        "ai_resp": "Validate ship-to ZIP, carrier coverage, delivery SLA, transit time, freight cost, warehouse selection, shipment split, ETA, and optimization tradeoffs.",
        "human_resp": "Review unserviceable ZIP exceptions, alternate delivery proposals, route constraints, or high freight cost exceptions.",
        "systems": "TMS, carrier coverage tables, freight rating engine, SLA rules, warehouse master, delivery calendars, optimization engine.",
        "exceptions": ["ZIP not serviceable", "Delivery SLA miss", "High freight cost", "Alternate route approval"],
        "module": "modules/logistics_validator.py",
        "master": "logistics-master-data.xlsx",
        "sheets": [
            ("Carrier_Coverage", "carrier, zip_prefix, serviceable, service_level, transit_days"),
            ("Freight_Rating", "carrier, zone, weight_min, weight_max, base_rate, per_kg_rate"),
            ("SLA_Rules", "scope_type, scope_id, service_level, max_transit_days, on_time_target_pct"),
            ("Warehouse_Master", "warehouse_id, name, zip, region, capacity, cutoff_time, carriers"),
            ("Delivery_Calendar", "warehouse_id, weekday, cutoff_time, ships, holiday_dates"),
        ],
        "how": ("Carrier coverage is checked for the ship-to ZIP prefix; the best ground carrier is selected, ETA is "
                "computed from transit time and compared to the customer SLA, and freight is rated on shipment weight. "
                "Unserviceable ZIPs and SLA breaches raise exceptions with alternatives."),
        "acs": [
            ("AC-01", "Validate delivery serviceability and ETA",
             "Given inventory is available for fulfillment\nWhen the AI agent evaluates the ship-to ZIP, carrier coverage, delivery SLA, delivery calendar, and transit time\nThen the system should determine whether the location is serviceable\nAnd calculate the estimated delivery date\nAnd proceed to fulfillment optimization if serviceable",
             "Serviceability and ETA are computed and shown; serviceable orders proceed to execution."),
            ("AC-02", "Recommend optimal fulfillment strategy",
             "Given the order is serviceable\nWhen the AI agent evaluates warehouse choice, inventory balance, freight cost, shipment split, and delivery SLA\nThen the system should recommend the optimal fulfillment option\nAnd display the selected warehouse or split plan, allocated quantities, freight cost, delivery ETA, and service rationale",
             "The selected warehouse, carrier, freight, ETA, and single/split plan are displayed."),
            ("AC-03", "Create exception when ZIP is not serviceable",
             "Given the ship-to ZIP is not serviceable by available carriers\nWhen the AI agent completes logistics validation\nThen the system should create a logistics exception\nAnd suggest alternatives such as nearest pickup location, alternate carrier, alternate ship-to, or revised delivery date\nAnd route the exception to the CSR for customer communication",
             "A ZIP_NOT_SERVICEABLE exception lists alternative pickup warehouses (sample: scenario-zip-not-serviceable.txt)."),
            ("AC-04", "Evaluate split shipment tradeoff",
             "Given the order can be fulfilled from multiple locations\nWhen the AI agent compares single-shipment and split-shipment options\nThen the system should evaluate freight cost, delivery SLA, inventory balance, warehouse capacity, and customer split-shipment preference\nAnd recommend the option with the best configured cost-service outcome",
             "The optimizer recommends single vs split shipment using freight/SLA and the customer split preference."),
        ],
        "scenarios": [
            ("ZIP not serviceable", "scenario-zip-not-serviceable.txt", "ZIP_NOT_SERVICEABLE + alternatives"),
        ],
    },
    "US-11": {
        "capability": "Exception-Based CSR Intervention and Governance",
        "layer": "Human-in-the-Loop Governance / Exception Handling",
        "persona": "CSR / Approver / Supervisor",
        "user_story": ("As a CSR, I want the AI agent to surface only exceptions that require human intervention, so that "
                       "routine orders are processed autonomously while complex B2B decisions remain governed."),
        "business_value": "Enables exception-based operations, reduces CSR workload, and maintains human control over high-risk decisions.",
        "ai_resp": "Detect exceptions, classify severity, route to the right role, and provide recommended action, rationale, and impacted order context.",
        "human_resp": "Approve, modify, reject, escalate, or comment on AI recommendations.",
        "systems": "Exception management queue, workflow engine, user roles, approval matrix, audit log, notification service.",
        "exceptions": ["All exception categories across intake, customer, product, pricing, approval, credit, inventory, logistics, compliance"],
        "module": "modules/exception_governance.py",
        "master": "exception-governance-master-data.xlsx",
        "sheets": [
            ("Severity_Matrix", "exception_type, severity, category, default_role"),
            ("Role_Routing", "category, role, escalation_role, queue"),
            ("SLA_Thresholds", "severity, sla_hours, escalation_hours, notify_channel"),
        ],
        "how": ("This cross-cutting layer inspects every stage result. With no exceptions it reports straight-through "
                "autonomous processing. When exceptions exist it classifies each by severity and category, routes to the "
                "responsible role with an escalation path, attaches resolution/escalation SLAs, and presents full decision "
                "context for the primary exception."),
        "acs": [
            ("AC-01", "Route exception with full decision context",
             "Given the AI agent identifies an exception during any stage\nWhen the exception is created\nThen the system should route it to the appropriate role\nAnd display the exception reason, impacted order lines, AI recommendation, business rule applied, confidence where available, and required action\nAnd pause automated progression until resolved",
             "Each exception is routed to a role/queue with severity and SLA; the primary exception shows full context."),
            ("AC-02", "Allow CSR to approve, modify, reject, or escalate recommendation",
             "Given an exception is assigned to a CSR or approver\nWhen the user reviews the AI recommendation\nThen the user should be able to approve, modify, reject, or escalate\nAnd the system should capture the user action, comments, timestamp, and final decision\nAnd resume automated orchestration when resolved",
             "The governance card states the available actions; resolution resumes orchestration (human-in-the-loop)."),
            ("AC-03", "Track exception SLA",
             "Given an exception is created\nWhen it remains unresolved beyond the configured SLA threshold\nThen the system should escalate to the next responsible role or supervisor\nAnd notify the assigned user with order number, exception type, age, and recommended next action",
             "SLA_Thresholds drive resolution and escalation timers with the configured notify channel per severity."),
        ],
        "scenarios": [
            ("Autonomous (no intervention)", "happy-autonomous.txt", "All clear — straight-through processing"),
            ("Governed exception", "scenario-governed-exception.txt", "Exception routed with severity, role, and SLA"),
        ],
    },
    "US-12": {
        "capability": "Order Creation, Downstream Fulfillment, Communication, and Audit Trail",
        "layer": "Execution Orchestration Layer",
        "persona": "CSR / Order Operations User / Customer",
        "user_story": ("As an order operations user, I want the AI agent to create downstream execution records and "
                       "communicate order confirmation details to the customer, so that the customer receives a confirmed "
                       "B2B order with price, ETA, fulfillment details, and required documentation."),
        "business_value": "Completes the autonomous PO-to-fulfillment journey and provides traceable confirmation to the customer.",
        "ai_resp": "Create ERP order, OMS request, WMS pick ticket, shipment order, confirmation, tracking update, documentation, and audit trail.",
        "human_resp": "Monitor exceptions, review failed integrations, and support customer follow-up.",
        "systems": "ERP, OMS, WMS, TMS, notification service, customer communication templates, document repository, audit log.",
        "exceptions": ["ERP creation failure", "OMS failure", "WMS failure", "Shipment creation failure", "Communication failure", "Audit gap"],
        "module": "modules/order_execution.py",
        "master": "execution-master-data.xlsx",
        "sheets": [
            ("Integration_Endpoints", "system, record_type, endpoint, mode, default_status, retry_limit"),
            ("Communication_Templates", "template_id, channel, subject, body_ref"),
            ("Document_Repository", "doc_id, doc_type, linked_to, file_ref"),
        ],
        "how": ("Once all validations pass, the agent creates mock ERP/OMS/WMS/TMS records and the customer confirmation "
                "(order number, PO, final price, contract, quantities, fulfillment source, carrier, ETA, payment terms, "
                "and compliance documents). A simulated downstream failure raises an execution exception and withholds the "
                "confirmation until resolved."),
        "acs": [
            ("AC-01", "Create downstream fulfillment records",
             "Given the order has passed all required validations and approvals\nWhen the AI agent completes orchestration\nThen the system should create the ERP sales order, OMS request, WMS pick ticket, and shipment order\nAnd store the final orchestration decision against the order record",
             "Mock ERP/OMS/WMS/TMS records are created with IDs and statuses (sample: happy-path.txt)."),
            ("AC-02", "Send B2B order confirmation to customer",
             "Given the order has been successfully created\nWhen the AI agent generates customer communication\nThen the customer should receive order acceptance confirmation\nAnd the communication should include confirmed order number, PO number, final price, contract reference, confirmed quantity, UOM, fulfillment source, ETA, tracking, payment terms, and required compliance documents",
             "The confirmation card consolidates all required fields from the upstream stages."),
            ("AC-03", "Maintain audit trail for AI and human decisions",
             "Given the AI agent processes a B2B purchase order\nWhen the order moves through orchestration\nThen the system should capture each automated decision, pricing rule, approval, exception, CSR intervention, fulfillment recommendation, and integration result\nAnd the audit trail should include timestamp, decision source, action, business rule, data source, and outcome",
             "Every stage emits an audit trail visible in the UI and aggregated across the pipeline."),
            ("AC-04", "Handle downstream integration failure",
             "Given the order is ready for downstream execution\nWhen ERP, OMS, WMS, TMS, or communication integration fails\nThen the system should create an execution exception\nAnd display the failed system, failure message, retry status, and recommended resolution\nAnd prevent customer confirmation until resolved or overridden",
             "An EXECUTION_FAILURE exception withholds confirmation and shows retry/resolution (sample: scenario-execution-failure.txt)."),
        ],
        "scenarios": [
            ("Successful execution", "happy-path.txt", "Order created + confirmation sent"),
            ("Downstream integration failure", "scenario-execution-failure.txt", "EXECUTION_FAILURE (mock WMS outage)"),
        ],
    },
}


# ── Document builder ──────────────────────────────────────────────────────────
def heading(doc, text, size=14, color=NAVY, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def kv_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    return t


def build_doc(story_id, s):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"{story_id} — {s['capability']}")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("PO Fulfillment AI Agent — Proof of Concept  |  Manager Summary")
    sr.italic = True
    sr.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared by: Sanjay Kumar Kesarvani    |    Date: 26 June 2026").font.size = Pt(10)

    heading(doc, "1. Story Overview")
    kv_table(doc, [
        ("Story ID", story_id),
        ("Epic", "Agentic Purchase Order-to-Fulfillment Orchestration for Enterprise B2B"),
        ("Layer", s["layer"]),
        ("Capability", s["capability"]),
        ("Persona", s["persona"]),
        ("User Story", s["user_story"]),
        ("Business Value", s["business_value"]),
        ("AI Responsibility", s["ai_resp"]),
        ("Human Responsibility", s["human_resp"]),
        ("Systems / Data Needed", s["systems"]),
        ("Exception Types Covered", "; ".join(s["exceptions"])),
    ])

    heading(doc, "2. Acceptance Criteria (story-wise) and how the POC satisfies them")
    for code, ac_title, gwt, poc in s["acs"]:
        heading(doc, f"{code} — {ac_title}", size=12, color=BLUE, space_before=8)
        for line in gwt.split("\n"):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(10)
        p = doc.add_paragraph()
        run = p.add_run("How the POC satisfies it: ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(poc)
        run2.font.size = Pt(10)

    heading(doc, "3. What was built in the POC")
    doc.add_paragraph(s["how"])
    kv_table(doc, [
        ("Module", s["module"]),
        ("Master data workbook", f"mock-data/{s['master']}"),
    ])
    heading(doc, "Master data sheets and fields", size=12, color=BLUE, space_before=6)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Sheet"
    hdr[1].text = "Fields (maximum possible captured for the POC)"
    for run in hdr[0].paragraphs[0].runs + hdr[1].paragraphs[0].runs:
        run.bold = True
    for name, fields in s["sheets"]:
        cells = t.add_row().cells
        cells[0].text = name
        cells[1].text = fields

    heading(doc, "4. Exception scenarios and sample files")
    doc.add_paragraph("Sample PO files for this story live in this same folder "
                      f"(sample-data/{story_id}/). Paste a file's contents into the app, "
                      "or open the app and submit it, to reproduce each outcome.")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Scenario"
    hdr[1].text = "Sample file"
    hdr[2].text = "Expected outcome"
    for c in hdr:
        for run in c.paragraphs[0].runs:
            run.bold = True
    for scen, fname, outcome in s["scenarios"]:
        cells = t.add_row().cells
        cells[0].text = scen
        cells[1].text = fname
        cells[2].text = outcome

    heading(doc, "5. How to demo")
    for step in [
        "Start the app:  python -m streamlit run app.py",
        "Open a sample file from this folder and copy its contents.",
        "Paste the PO text into the chat box (or upload an Excel PO) and submit.",
        "Watch the AI agent process each stage; this story's stage is "
        f"'{s['capability']}'.",
        "For exception files, the agent pauses and shows the exception with decision context; "
        "for the happy path it passes and continues to the next stage.",
    ]:
        p = doc.add_paragraph(step, style="List Number")
        for run in p.runs:
            run.font.size = Pt(10)

    note = doc.add_paragraph()
    nr = note.add_run("Note: This is a local POC. There is no real ERP/CRM/WMS/OMS/TMS/SMTP — "
                      "every integration point is simulated with the mock master-data workbooks above.")
    nr.italic = True
    nr.font.size = Pt(9)

    out = os.path.join(BASE, story_id, f"{story_id}_Summary_for_Manager.docx")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print("Created:", os.path.join("sample-data", story_id, f"{story_id}_Summary_for_Manager.docx"))


def main():
    for sid, s in STORIES.items():
        build_doc(sid, s)
    print(f"\n{len(STORIES)} manager documents created.")


if __name__ == "__main__":
    main()
