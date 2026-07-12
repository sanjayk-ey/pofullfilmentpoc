"""
create_master_data_reference.py
Generates docs/Master_Data_Reference.docx — a business-friendly reference guide
listing every master data workbook, every sheet inside it, and the plain-English
use case each sheet serves.

Run:  python create_master_data_reference.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "docs")
os.makedirs(OUT_DIR, exist_ok=True)

NAVY  = RGBColor(0x1F, 0x2A, 0x44)
BLUE  = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x1E, 0x7D, 0x44)
GREY_BG   = "F3F4F6"
HEADER_BG = "1F2A44"
BAND_BG   = "F8FAFC"

# ─── Master-data catalog ──────────────────────────────────────────────────────
# One entry per workbook. Each `sheets` list is in the same order the sheet
# appears inside the workbook. `use_case` explains, in plain words, what the
# validators actually use that sheet for at runtime.
CATALOG = [
    ("customer-master-data.xlsx",
     "US-02 Account Hierarchy & Ship-To Validation",
     [
        ("Customer_Master",
         "The single source of truth for who each customer is. Holds ERP ID, CRM "
         "ID, legal name, tax ID, primary contact, credit rating, and status. Used "
         "at account validation to resolve the customer named on the PO to a "
         "unique internal account (and to raise UNMATCHED or DUPLICATE customer "
         "exceptions)."),
        ("Account_Hierarchy",
         "Maps every branch to its regional division and its global parent. Used "
         "to walk the account tree upwards from the ship-to and confirm the "
         "ship-to sits under the same global parent as the ordering customer "
         "(HIERARCHY_MISMATCH check)."),
        ("Ship_To_Master",
         "All valid ship-to locations with their ZIP, address, branch, active "
         "flag, split/backorder tolerance, and default delivery instructions "
         "(dock hours, contact phone, access notes). Used to look up the "
         "ship-to ZIP from the PO and confirm it is registered "
         "(INVALID_SHIP_TO check). The `default_delivery_instructions` "
         "column captures the site's routine handling rules; the UI shows "
         "delivery instructions from the PO only (per-transaction) rather "
         "than pulling this default in as a fallback."),
        ("Hierarchy_Rules",
         "Which rules apply at which level of the hierarchy — pricing tier, "
         "product visibility, budget limit, approval routing, and fulfillment "
         "rule ID. The most specific level (ship-to > branch > regional > "
         "global) wins."),
        ("Fulfillment_Rules",
         "The actual business content behind each fulfillment rule ID: preferred "
         "warehouse, alternate warehouses, restricted warehouses, split allowed, "
         "backorder allowed, max backorder days, min order qty, delivery SLA "
         "days, allocation priority. Used by US-09 Inventory and US-10 "
         "Fulfillment Optimization stages."),
     ]),

    ("buyer-master-data.xlsx",
     "US-03 Buyer Authorization & Product Visibility",
     [
        ("Buyer_Profiles",
         "The master list of authorized buyers (contact persons on POs). Contains "
         "buyer ID, name, email, employer, status. Used to verify that the person "
         "named on the PO is actually authorized to place orders on that "
         "customer's behalf (UNAUTHORIZED_BUYER check)."),
        ("User_Permissions",
         "Per-buyer permission flags — order value limit, product categories "
         "they may order, geographies they may ship to. Enforced by the buyer "
         "authorization stage."),
        ("Cost_Centers",
         "Master list of cost centers with the branch each cost center belongs "
         "to. Used to validate that the cost center on the PO is real and "
         "belongs to the buyer's branch (INVALID_COST_CENTER check)."),
        ("Product_Visibility_Rules",
         "Which product categories a buyer's role/profile can order — enforces "
         "corporate contracts (e.g., 'Facilities buyers may not order raw "
         "chemicals'). Feeds RESTRICTED_PRODUCT exceptions."),
     ]),

    ("product-master-data.xlsx",
     "US-04 Product Matching & UOM Conversion",
     [
        ("Product_Master",
         "SKU master — name, description, category, status (active / obsolete / "
         "discontinued), unit list price, weight, hazardous flag. Used to "
         "validate each ordered SKU exists and is still sellable "
         "(OBSOLETE_SKU / PRODUCT_CONFIG_EXCEPTION)."),
        ("Product_Attributes",
         "Extended attributes such as finish, size, colour, spec. Used to render "
         "meaningful line-item descriptions and to power substitution logic."),
        ("UOM_Conversions",
         "Conversion factors between units of measure (e.g., 1 CASE = 24 EA, "
         "1 PALLET = 40 CASE). Used to normalize the customer's UOM to the "
         "vendor's stocking UOM (INVALID_UOM check when a mapping is missing)."),
        ("Substitution_Rules",
         "When an ordered SKU is obsolete or unavailable, this table names the "
         "recommended substitute. Powers the 'Recommend substitute SV-220' "
         "action shown in the Product Match stage."),
        ("Compatibility_Rules",
         "Which SKUs cannot be ordered together and which SKUs must be ordered "
         "together (e.g., valve + cartridge kit). Used to protect the customer "
         "from placing an unusable order."),
     ]),

    ("compliance-master-data.xlsx",
     "US-05 Compliance & Regulatory Checks",
     [
        ("Compliance_Rules",
         "The rulebook: which product categories are hazardous, which require "
         "SDS attachment, which require export licences. Used to decide what "
         "checks each order line must pass."),
        ("Regional_Restrictions",
         "Which SKUs cannot ship to which regions or countries (Prop-65, state "
         "chemical bans, export controls). Feeds COMPLIANCE_RESTRICTION "
         "exceptions."),
        ("SDS_Repository",
         "Safety Data Sheet availability per SKU. Used to raise MISSING_SDS when "
         "a hazardous SKU is missing its current SDS document."),
        ("Product_Eligibility",
         "Whether a specific SKU is eligible for sale in a specific country / "
         "region / customer channel. Enforces distribution agreements."),
     ]),

    ("pricing-master-data.xlsx",
     "US-06 Pricing & Discount Waterfall",
     [
        ("Price_List",
         "The base list price for every SKU. The starting point for the pricing "
         "waterfall (the 'sticker price' before any discounts)."),
        ("Contracts",
         "Customer-specific contract prices with validity dates. Overrides the "
         "list price during the contract term. Expired contracts fall back to "
         "list price (see US-06 expired-contract scenario)."),
        ("Volume_Tiers",
         "Quantity-based discount tiers (e.g., 100-499 = 3% off, 500-999 = 5%, "
         "1000+ = 7%). Applied on top of contract price."),
        ("Rebates",
         "Post-order rebate percentages by customer or period. Shown as line 5 "
         "of the price waterfall; does not change the invoice price but affects "
         "margin analytics."),
        ("Promotions",
         "Time-bound promotional discounts (e.g., Q3 kitchen faucet promotion, "
         "5% off through 30-Sep). Automatically expires."),
        ("Surcharges",
         "Location-based, hazmat, or fuel surcharges applied to the freight "
         "line (e.g., +$25 remote-area surcharge for AK/HI ZIP codes)."),
        ("Freight_Terms",
         "Freight allowance rules per customer — FOB destination, prepaid & "
         "add, freight-included, etc. Includes base freight, minimum freight, "
         "and per-KG rate so the AI can calculate the shipping charge for "
         "every order."),
        ("Tax_Rates",
         "State-level (and country-level) sales tax / VAT rates keyed by "
         "region code. The pricing engine looks up the ship-to state and "
         "applies the correct tax percentage on (subtotal + surcharges); the "
         "resulting tax amount is shown on-screen and stored on the order."),
        ("Margin_Policy",
         "The guardrails: minimum margin floor per product category. If the "
         "discount stack would push margin below the floor, the pricing engine "
         "raises PRICING_EXCEPTION and asks CSR to approve or reject."),
        ("Raw_Material_Index",
         "Current raw material cost index (copper, brass, PVC). Feeds cost-plus "
         "recalculation of the margin floor when raw material costs move."),
     ]),

    ("budget-master-data.xlsx",
     "US-07 Budget Availability & Approval Routing",
     [
        ("Budget_Master",
         "Approved budget per cost center or department for the current period, "
         "plus remaining balance. Used to check the order does not overrun the "
         "budget (BUDGET_EXCEEDED exception)."),
        ("Cost_Centers",
         "Cost center master used for lookups (mirrors the copy in "
         "buyer-master-data). Confirms the payer entity exists and is active."),
        ("Approval_Matrix",
         "Order-value thresholds and who approves at each tier "
         "(e.g., < $10k self-approve, $10k-$100k branch manager, "
         "$100k-$1M regional director, > $1M CFO). Drives the "
         "APPROVAL_REQUIRED routing."),
        ("Buyer_Authority",
         "Each buyer's individual self-approval limit — may be tighter than "
         "the default tier for that buyer's cost center. The tighter of the "
         "two wins."),
        ("Approvers",
         "Names and email addresses of approvers per role and level. Used to "
         "populate the 'Approval email sent to' line and audit trail."),
     ]),

    ("credit-master-data.xlsx",
     "US-08 Credit Check & Payment Terms",
     [
        ("Credit_Master",
         "Credit limit and available credit per customer account. Used to check "
         "the new order does not push the customer over their limit "
         "(CREDIT_HOLD exception)."),
        ("Invoice_Aging",
         "Overdue invoice buckets per customer (0-30, 31-60, 61-90, 90+ days). "
         "Aged debt reduces available credit and may trigger a hold even when "
         "the raw limit still has room."),
        ("Payment_History",
         "Historical payment behaviour per customer — on-time rate, average "
         "days-to-pay, recent chargebacks. Signals risk trends to the credit "
         "analyst."),
        ("Payment_Terms",
         "Standard payment terms per customer (Net 15 / Net 30 / Net 45 / Net "
         "60 / Prepaid / Letter of Credit). Shown on the customer order "
         "confirmation and used by the invoicing system."),
        ("Risk_Signals",
         "Watch-list customers, industry-wide risk flags, fraud signals. "
         "Overrides otherwise-clean credit and routes the order to the "
         "credit team for manual review."),
     ]),

    ("inventory-master-data.xlsx",
     "US-09 Inventory, ATP & Allocation",
     [
        ("Plant_Stock",
         "Stock at manufacturing plants (upstream supply). Not directly "
         "shippable to customers but feeds ATP for future replenishment "
         "windows."),
        ("DC_Stock",
         "Distribution centre on-hand quantity and quantity already reserved "
         "for other demand. Allocatable stock = on-hand minus reserved. The "
         "single most important table for the inventory stage."),
        ("In_Transit",
         "Inventory currently in transit between plants and DCs, with expected "
         "arrival date. Added to ATP for orders whose requested-delivery-date "
         "is on or after the in-transit ETA."),
        ("ATP",
         "Available-to-promise per SKU — the aggregated 'what can we actually "
         "commit right now' view combining on-hand, in-transit, and reserved. "
         "Cited in the sourcing plan shown to the customer."),
        ("Allocation_Rules",
         "When stock is constrained and multiple orders compete for the same "
         "SKU, this table decides who wins by customer tier (GOLD > SILVER > "
         "BRONZE). Drives ALLOCATION_CONFLICT decisions."),
        ("Fulfillment_Preferences",
         "Legacy customer-level fulfillment overrides. Modern rules live in "
         "customer-master-data.Fulfillment_Rules; this table remains for "
         "backwards compatibility."),
     ]),

    ("logistics-master-data.xlsx",
     "US-10 Fulfillment Optimization, Logistics & Delivery SLA",
     [
        ("Carrier_Coverage",
         "Which carriers service which ZIP prefixes with what service level "
         "(GROUND / EXPRESS) and transit days. Used first to confirm the "
         "ship-to ZIP is serviceable at all (ZIP_NOT_SERVICEABLE exception)."),
        ("Freight_Rating",
         "Freight cost per carrier per zone per weight band (base rate + per-kg "
         "rate). The optimization stage rates every candidate plan (preferred "
         "DC, alternate DC, split) against this table and picks the cheapest "
         "feasible plan."),
        ("SLA_Rules",
         "Delivery SLA commitments per customer tier — maximum transit days "
         "and on-time target percentage. Any candidate plan whose transit "
         "exceeds the SLA is rejected."),
        ("Warehouse_Master",
         "Master list of warehouses / DCs — name, ZIP, region, capacity, "
         "cutoff time, supported carriers, hazmat capability. Used to compute "
         "origin-to-destination zones for freight rating."),
        ("Delivery_Calendar",
         "Ship days, cutoff times, and holiday closures per warehouse. Used "
         "to compute a realistic ETA (skipping weekends and holidays)."),
     ]),

    ("exception-governance-master-data.xlsx",
     "US-11 Exception Governance & Routing",
     [
        ("Severity_Matrix",
         "Every exception type mapped to a severity band — CRITICAL / HIGH / "
         "MEDIUM / LOW. Drives colour coding, alerting, and escalation."),
        ("Role_Routing",
         "Which human role owns each exception type (e.g., CREDIT_HOLD -> "
         "Credit Analyst, COMPLIANCE_RESTRICTION -> Compliance Officer, "
         "PRICING_EXCEPTION -> Pricing Manager, INVENTORY_SHORTAGE -> CSR). "
         "Determines who receives the notification."),
        ("SLA_Thresholds",
         "Response-time SLA per role per severity (e.g., a CRITICAL exception "
         "must be responded to within 30 min; LOW within 2 business days). "
         "Feeds escalation timers."),
     ]),

    ("execution-master-data.xlsx",
     "US-12 Order Execution & Confirmation",
     [
        ("Integration_Endpoints",
         "Downstream system URLs (ERP, OMS, WMS, TMS, SMTP) that the "
         "execution stage calls to create the sales order, pick ticket, "
         "shipment, and send the customer confirmation email."),
        ("Communication_Templates",
         "The email and document templates used to notify customers and "
         "internal teams — order acknowledgement, shipment notification with "
         "tracking, approval request, exception alert."),
        ("Document_Repository",
         "Where generated documents are stored — order confirmations, packing "
         "slips, invoices, proofs of delivery. Referenced from the audit trail."),
     ]),
]


# ─── Doc-writing helpers ──────────────────────────────────────────────────────
def _shade(cell, hex_color):
    """Apply a background fill colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_col_widths(table, widths_in_inches):
    """Force column widths — Word ignores unless applied to every cell."""
    for row in table.rows:
        for i, w in enumerate(widths_in_inches):
            row.cells[i].width = Inches(w)


def main():
    doc = Document()
    # Landscape orientation gives the "use case" column enough room to breathe
    section = doc.sections[0]
    section.orientation = 1  # WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width  = new_w
    section.page_height = new_h
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("PO Fulfilment Order Assistant")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("Master Data Reference")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    total_sheets = sum(len(sheets) for _, _, sheets in CATALOG)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"{len(CATALOG)} master workbooks   |   {total_sheets} tables   |   "
        f"the complete reference data behind every decision stage"
    ).italic = True

    doc.add_paragraph()

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run(
        "This document lists every master-data Excel workbook used by the AI "
        "orchestration agent, every table (sheet) inside it, and the plain-English "
        "use case each table serves at runtime. Workbooks are ordered by the user "
        "story that consumes them (US-02 through US-12). All files live under "
    )
    intro.add_run("mock-data/").font.name = "Consolas"
    intro.add_run(" in the project root.")

    doc.add_paragraph()

    # ─── The main reference table ───────────────────────────────────────────
    n_rows = 1 + total_sheets  # header + one row per sheet
    table = doc.add_table(rows=n_rows, cols=3)
    table.style = "Light Grid Accent 1"
    table.autofit = False

    widths = [1.9, 1.9, 6.7]  # inches (landscape total ~10.5" usable)

    # Header row
    hdr = table.rows[0]
    for i, txt in enumerate(("Master file name", "Tables in each master file",
                             "Use case of each table")):
        c = hdr.cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
        _shade(c, HEADER_BG)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    row_idx = 1
    for wb_name, wb_purpose, sheets in CATALOG:
        for j, (sheet_name, use_case) in enumerate(sheets):
            row = table.rows[row_idx]

            # Column 1 — workbook name (only on the first row of each workbook,
            # merged vertically with subsequent rows below).
            if j == 0:
                c1 = row.cells[0]
                c1.text = ""
                p = c1.paragraphs[0]
                r1 = p.add_run(wb_name + "\n")
                r1.bold = True
                r1.font.size = Pt(10)
                r1.font.color.rgb = NAVY
                r2 = p.add_run(wb_purpose)
                r2.italic = True
                r2.font.size = Pt(9)
                r2.font.color.rgb = GREEN

            # Column 2 — sheet name
            c2 = row.cells[1]
            c2.text = ""
            r2 = c2.paragraphs[0].add_run(sheet_name)
            r2.bold = True
            r2.font.name = "Consolas"
            r2.font.size = Pt(10)

            # Column 3 — use case
            c3 = row.cells[2]
            c3.text = use_case

            # Alternate band shading across workbooks
            if list(CATALOG).index((wb_name, wb_purpose, sheets)) % 2 == 1:
                for c in row.cells:
                    _shade(c, BAND_BG)

            row_idx += 1

        # Vertically merge column 1 across all rows of this workbook
        top = 1 + row_idx - len(sheets)
        bot = row_idx - 1
        if bot > top:
            table.cell(top, 0).merge(table.cell(bot, 0))

    _set_col_widths(table, widths)

    # ─── Footer notes ───────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Notes:  ").bold = True
    note.add_run(
        "(1) All workbooks are regenerated in one shot by running "
    )
    note.add_run("python create_master_data.py").font.name = "Consolas"
    note.add_run(
        " and, for customer-master-data.xlsx only, "
    )
    note.add_run("python create_customer_master_data_excel.py").font.name = "Consolas"
    note.add_run(
        ".  (2) Sheets are loaded on first use by the validator that owns them "
        "and cached in-process for the life of the Streamlit server — changes "
        "to a workbook require an app restart to take effect."
    )

    out = os.path.join(OUT_DIR, "Master_Data_Reference.docx")
    doc.save(out)
    print(f"Created: {out}  ({len(CATALOG)} workbooks, {total_sheets} tables)")


if __name__ == "__main__":
    main()
