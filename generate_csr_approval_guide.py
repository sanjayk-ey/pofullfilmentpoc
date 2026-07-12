"""Generate 'CSR Approval Conditions & Test Guide' (.docx).

This is a REVIEW document. It enumerates every condition in the Autonomous
PO-to-Fulfillment Orchestration where the Order assistant pauses for a Customer Service
Representative (CSR) decision, explains the process for each, and gives exact,
copy-paste test steps using the real mock master-data values.

Source of truth for the decision layers:
    Autonomous_PO_to_Fulfillment_Orchestration.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── EY palette ────────────────────────────────────────────────────────────────
EY_YELLOW = RGBColor(0xFF, 0xE6, 0x00)
EY_DARK   = RGBColor(0x2E, 0x2E, 0x38)
HDR_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
GREY      = RGBColor(0x59, 0x59, 0x59)
AMBER     = RGBColor(0xB0, 0x60, 0x00)
GREEN     = RGBColor(0x1E, 0x7A, 0x34)
RED       = RGBColor(0xB0, 0x1C, 0x1C)
HDR_FILL  = "1F4E79"
ZEBRA     = "F2F5FA"


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell(cell, text, bold=False, size=9, color=None, fill=None, italic=False):
    if fill:
        shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    for i, line in enumerate(str(text).split("\n")):
        run = (p if i == 0 else cell.add_paragraph()).add_run(line)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color


def heading(doc, text, size=15, color=HDR_BLUE, before=12, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def subheading(doc, text, size=12, color=EY_DARK, before=8, after=2):
    return heading(doc, text, size=size, color=color, before=before, after=after)


def para(doc, text, size=10, color=None, bold=False, italic=False, before=0, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        bold = False
        text = it
        if isinstance(it, tuple):
            lead, text = it
            r = p.add_run(lead + " ")
            r.bold = True
            r.font.size = Pt(size)
        r = p.add_run(text)
        r.font.size = Pt(size)


def table(doc, headers, rows, widths=None, header_fill=HDR_FILL, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, bold=True, size=9,
                 color=RGBColor(0xFF, 0xFF, 0xFF), fill=header_fill)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        fill = ZEBRA if (zebra and i % 2 == 0) else None
        for j, val in enumerate(row):
            set_cell(cells[j], val, size=9, fill=fill)
    return t


def status_badge(doc, kind):
    """kind: 'live' or 'new'"""
    if kind == "live":
        para(doc, "Status: LIVE today", size=9, color=GREEN, bold=True, after=2)
    else:
        para(doc, "Status: ENHANCED in this update", size=9, color=AMBER, bold=True, after=2)


# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Title ─────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
tr = t.add_run("CSR Approval Conditions & Test Guide")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = EY_DARK
s = doc.add_paragraph()
sr = s.add_run("Autonomous Purchase Order-to-Fulfillment Orchestration  |  Human-in-the-Loop Governance")
sr.bold = True
sr.font.size = Pt(13)
sr.font.color.rgb = GREY
para(doc, "Prepared for review. Lists every point in the workflow where the AI "
          "agent pauses and asks a Customer Service Representative (CSR) to decide, "
          "what the CSR sees, and exactly how to test each one.",
     size=10, color=GREY, italic=True, before=4)

# ── 1. How human-in-the-loop works ────────────────────────────────────────────
heading(doc, "1. How Human-in-the-Loop Governance Works")
para(doc, "The Order assistant processes each purchase order autonomously through a "
          "sequence of decision layers. Whenever a business rule is breached, data "
          "is missing, or data is ambiguous, the agent stops, explains what it "
          "found, and hands the decision to a CSR. The order does not advance until "
          "the CSR responds. Depending on the situation the CSR can:")
bullets(doc, [
    ("Approve", "accept the AI's recommendation and continue processing."),
    ("Modify / Enter", "provide a correction (e.g. the right SKU, quantity, or address)."),
    ("Reject", "stop the order (used only where overriding is a valid business choice)."),
    ("Escalate", "route the exception to the responsible team (Pricing, Finance, Duplicate-PO, etc.)."),
])
para(doc, "There are two families of CSR touch-points:")
bullets(doc, [
    ("Intake resolution", "before the pipeline runs, the agent cleans up the raw PO "
     "(missing buyer, partial ship-to, wrong/missing SKU, bad quantity, duplicate PO)."),
    ("Decision-layer approvals", "inside the pipeline, each layer can raise an "
     "exception that needs CSR sign-off (substitution, pricing, credit, inventory, logistics)."),
])

# ── 2. Decision layers & planned changes ──────────────────────────────────────
heading(doc, "2. Decision Layers (aligned to the architecture document)")
para(doc, "Headings will be renamed to match the nine decision layers in the "
          "architecture document. The mapping and the CSR gate for each layer:")
table(doc,
    ["Decision layer (new heading)", "Current app stage", "CSR approval in this layer?"],
    [
        ["Intake", "PO extraction & intake resolver", "Yes - missing/ambiguous data"],
        ["Customer Validation", "Corporate Account Hierarchy & Ship-To (+ Buyer Authorization folded in)", "Strategic-customer review (soft)"],
        ["Product Match", "Product Matching + UOM (+ Compliance folded in)", "Yes - obsolete/inactive SKU substitution"],
        ["Pricing and Promo", "Enterprise B2B Pricing Engine", "Yes - discount > policy / margin below floor"],
        ["Approval", "Budget, Spend Limit & Approval Routing", "Auto-approve or route to approver"],
        ["Credit", "Credit, Payment Terms & Financial Risk", "Yes - credit hold override / send to Finance"],
        ["Inventory Checks", "Inventory Availability & ATP", "Yes - shortage / partial-shipment proposal"],
        ["Logistics", "Fulfillment, Logistics & Delivery SLA (Optimization inside)", "Yes - ZIP not serviceable / alternate"],
        ["Optimization", "(shown inside Logistics: Option A/B/C cost comparison)", "No (informational)"],
    ],
    header_fill=HDR_FILL)

subheading(doc, "2.1  Changes being made in this update")
bullets(doc, [
    ("Resolved Account Hierarchy", "will additionally display Customer Tier, "
     "Payment Terms, and Distributor/Contractor classification."),
    ("Customer master data", "gains distributor-classification fields and a "
     "buying-history summary; buying history is checked during Customer Validation "
     "and Product Match."),
    ("Product Match buttons", "\"Approve suggestion\" is renamed to \"Approve\"; "
     "the \"Reject\" button is removed for the obsolete-product substitution."),
    ("Stage headings", "renamed to the nine decision-layer names above."),
    ("Pricing and Promo", "becomes an interactive CSR gate (Approve exception / "
     "Reject / Escalate) instead of an automatic e-mail hold."),
])
para(doc, "Note: 'LIVE today' below means the behavior already works; 'ENHANCED in "
          "this update' means it changes as part of this request.",
     size=9, color=GREY, italic=True)

# ── 3. Catalog of CSR-approval conditions ─────────────────────────────────────
heading(doc, "3. Catalog of CSR-Approval Conditions")

conditions = [
    # (num, title, layer, status, trigger, ai_shows, options, test_steps)
    ("3.1", "Missing or unresolved buyer", "Intake", "live",
     "The PO has no buyer name, or the name/email cannot be matched to the buyer master.",
     "The agent lists close matches from the buyer master (name, role, customer) and "
     "offers a field to type the correct buyer.",
     [("Approve / Pick", "use the suggested buyer."),
      ("Enter", "type the correct buyer name/email."),
      ("Escalate", "route to data stewardship.")],
     ["Use CSR-Approval-PO (buyer email procurement@greatlakesps.com is not in the "
      "buyer master), or change the email in Happy-Flow-PO to unknown@nowhere.com.",
      "At Intake the agent flags the unresolved buyer and shows suggestions.",
      "Pick a suggestion OR type john.miller@glps.com and click Approve; "
      "confirm processing resumes step by step."]),

    ("3.2", "Partial or missing ship-to address", "Intake", "live",
     "Ship-to has only a company/branch name (no street/ZIP), or the address is incomplete.",
     "The agent looks up the registered ship-to for that customer in the ship-to "
     "master and proposes the full address; a field allows entering a different address.",
     [("Approve", "use the address found in master data."),
      ("Enter", "type a different/corrected address."),
      ("Escalate", "route to logistics data team.")],
     ["Use CSR-Approval-PO (its SHIP TO has only 'Great Lakes - Ketchikan Project "
      "Site', no street).",
      "At Intake the agent proposes the full registered address (Ketchikan, ZIP 99950).",
      "Click Approve to accept, or type an alternative address, then continue."]),

    ("3.3", "Missing SKU (description + qty only)", "Intake", "live",
     "An order line has a description and quantity but no product code.",
     "The agent matches the description to the product master and shows the matched "
     "SKU with a confidence score; a field allows typing the correct SKU.",
     [("Approve", "accept the matched SKU."),
      ("Enter", "type the correct SKU (validated on the spot)."),
      ("Escalate", "route to product data team.")],
     ["Use CSR-Approval-PO line 3 ('Tank-to-Bowl Gasket Kit', qty 50, no code).",
      "The agent matches it to SKU-SEL-1150.",
      "Click Approve, or type a SKU; invalid/duplicate SKUs are rejected instantly.",
      "Confirm the pipeline runs fully after the decision."]),

    ("3.4", "Wrong / invalid SKU", "Intake", "live",
     "A product code is present but does not exist in the product master.",
     "The agent reports the code is not recognized and asks the CSR to enter a valid "
     "SKU; entries are validated immediately (must exist, must not duplicate an existing line).",
     [("Enter", "type the correct SKU (must pass on-the-spot validation)."),
      ("Escalate", "route to product data team.")],
     ["Use CSR-Approval-PO line 2 ('PN-DRAIN-STD') - a non-standard code.",
      "The agent matches by description to SKU-DRN-3010 (or asks for a valid SKU).",
      "Type an invalid code (e.g. SKU-XXX-0000) to see the instant error; then a "
      "duplicate of an existing line to see the duplicate error.",
      "Enter a valid SKU and confirm the process continues."]),

    ("3.5", "Invalid / zero quantity", "Intake", "live",
     "An order line quantity is 0, blank, or not a positive number.",
     "The agent flags the line and asks the CSR to enter a valid quantity for that SKU.",
     [("Enter", "type a positive quantity."),
      ("Escalate", "route to CSR supervisor.")],
     ["Use CSR-Approval-PO line 4 (SKU-VLV-2201, Qty = 0), or edit any PO line to Qty = 0.",
      "At Intake the agent asks for a corrected quantity.",
      "Type a positive number (e.g. 25) and confirm processing resumes."]),

    ("3.6", "Duplicate purchase order", "Intake", "live",
     "The same PO number for the same customer has already been submitted (source-of-truth check).",
     "The agent auto-rejects the order and shows: 'This order will not be processed.' "
     "No approve/override is offered - a duplicate cannot be re-processed.",
     [("Escalate", "send to the Duplicate-PO team to reconcile against the original. "
       "This is the ONLY available action.")],
     ["Submit Happy-Flow-PO (PO-2026-30001) once so it is recorded.",
      "Submit Happy-Flow-PO again -> the agent detects the duplicate and auto-rejects.",
      "Confirm only the Escalate button is shown (no Approve/Reject).",
      "(To reset between demos, clear data/submitted_pos.json.)"]),

    ("3.7", "Obsolete / inactive SKU - substitution", "Product Match", "new",
     "An ordered SKU has status OBSOLETE or INACTIVE and an approved substitute exists.",
     "The agent shows original SKU, recommended substitute, compatibility, price "
     "impact, availability, and rationale, then asks: 'Approve substitution?'",
     [("Approve", "replace with the recommended substitute and continue. "
       "(Button renamed from 'Approve suggestion' to 'Approve'.)"),
      ("Modify / Enter", "type a different substitute SKU."),
      ("Escalate", "route to product management. (The 'Reject' button is removed here.)")],
     ["Use CSR-Approval-PO line 1: SKU-CTG-1000 (OBSOLETE) qty 40.",
      "The agent recommends SKU-CTG-4520 (FULL compatibility, +13.6% price, in stock).",
      "Confirm the buttons are 'Approve', 'Modify/Enter', and 'Escalate' (no Reject).",
      "Click Approve and confirm the substituted line flows through the rest of the pipeline.",
      "Alt: SKU-VLV-2000 (INACTIVE) -> substitute SKU-VLV-2201."]),

    ("3.8", "Pricing exception - discount / margin", "Pricing and Promo", "new",
     "Effective discount vs list price exceeds the family policy cap (the '>10%' rule) "
     "OR the resulting margin is below the family floor.",
     "The agent shows List / Contract / Volume-discount / Final price and the margin "
     "impact, then asks: 'Requested discount exceeds policy. Margin impact = $X. "
     "Approve exception?' (Interactive CSR gate - replaces the old e-mail hold.)",
     [("Approve exception", "accept the pricing and continue."),
      ("Reject", "decline the pricing (order stops)."),
      ("Escalate", "route to the Pricing Manager/Director per the margin policy.")],
     ["Paste a PO for Great Lakes Plumbing (CUST-1001) with line SKU-CTG-4520 qty "
      "2000 (volume tier 10% + rebates 3% pushes effective discount ~25% > 18% cap).",
      "Other quick breaches: SKU-DRN-3010 qty 100 (~16.7% > 15%); "
      "SKU-VLV-2201 qty 50 (~15.5% > 12%).",
      "The Pricing and Promo layer pauses with the discount/margin message.",
      "Click Approve exception -> continues; try Escalate -> routed to Pricing Manager."]),

    ("3.9", "Credit hold", "Credit", "live",
     "Customer is on credit HOLD, is over available credit, has overdue invoices, "
     "or is flagged for fraud/watchlist.",
     "The agent shows credit limit, available credit, overdue amount and risk, then "
     "asks: 'Customer exceeded credit limit. Approve override or send to Finance?'",
     [("Approve override", "accept the credit risk and continue."),
      ("Escalate", "send to Finance for review.")],
     ["Simplest: submit the CSR-Approval-PO demo file - its ~$38.6K order for "
      "CUST-1001 (Great Lakes) exceeds the $30K available credit and pauses here.",
      "Or paste a PO that resolves to CUST-7000 (Midtown Building Supply) - status "
      "HOLD, $0 available, $200K overdue, fraud flag; or CUST-1002 (Eastern Kitchen "
      "& Bath) - $155K overdue.",
      "The Credit layer pauses with the credit-hold message.",
      "Click Approve override -> continues; or Escalate -> routed to Finance."]),

    ("3.10", "Inventory shortage - partial shipment", "Inventory Checks", "live",
     "Requested quantity is not fully available across plants/DCs within lead time.",
     "The agent shows on-hand by location and proposes a split (e.g. '400 now, 100 "
     "next week') so the customer proposal can be approved.",
     [("Approve", "accept the partial/split proposal and continue."),
      ("Escalate", "route to supply planning.")],
     ["Paste a PO for CUST-1001, ship-to ZIP 60639, line SKU-SHS-7700 qty 10 "
      "(only ~8 on hand at Chicago DC; ATP 5).",
      "The Inventory Checks layer pauses and proposes a partial/backorder split.",
      "Click Approve to accept the proposal; confirm processing continues."]),

    ("3.11", "ZIP not serviceable / delivery constraint", "Logistics", "live",
     "No carrier serviceable coverage exists for the ship-to ZIP prefix.",
     "The agent reports the ZIP is not serviceable and suggests the nearest pickup / "
     "alternate arrangement for the customer.",
     [("Approve", "accept the alternate arrangement and continue."),
      ("Escalate", "route to logistics.")],
     ["Paste a PO for CUST-1001 with ship-to ZIP 99950 (Ketchikan, AK - registered "
      "ship-to ST-AK-006; prefix 999 is not serviceable).",
      "The Logistics layer pauses with the not-serviceable message and alternate.",
      "Click Approve to accept the alternate; or Escalate."]),

    ("3.12", "Approval routing (spend / self-approval limit)", "Approval", "live",
     "Order value exceeds the buyer's self-approval limit, or the buyer cannot self-approve.",
     "The agent evaluates margin, discount limits and the approval matrix; if within "
     "policy it auto-approves, otherwise it routes to the mapped approver.",
     [("(Automatic)", "auto-approve when within policy (e.g. value < limit, margin ok)."),
      ("Route to approver", "sent to Branch/Regional/Corporate approver per the matrix.")],
     ["Happy-Flow-PO (~$3K, buyer John Miller, self-approve limit $250K) -> auto-approved.",
      "To force routing: use buyer Linda Park (linda.park@glps.com, cannot self-approve) "
      "or an order total above $250,000 for John Miller.",
      "Confirm the Approval layer shows auto-approval vs routed-to-approver."]),

    ("3.13", "Strategic customer review", "Customer Validation", "new",
     "The resolved customer is a Strategic-tier account (per Customer Tier in master data).",
     "During Customer Validation the agent surfaces the tier, distributor "
     "classification, payment terms and buying-history summary for CSR awareness "
     "(soft review trigger from the architecture document).",
     [("Acknowledge / Continue", "CSR reviews the strategic-account context and proceeds."),
      ("Escalate", "route to account management if special handling is needed.")],
     ["Use any PO for CUST-1001 (Great Lakes Plumbing - tier 'Strategic').",
      "Confirm the Resolved Account Hierarchy shows Customer Tier = Strategic, "
      "Payment Terms, and Distributor classification.",
      "Compare with CUST-5001 (Standard) to see the difference."]),
]

for num, title, layer, status, trigger, ai_shows, options, steps in conditions:
    subheading(doc, f"{num}  {title}")
    para(doc, f"Decision layer: {layer}", size=9, color=HDR_BLUE, bold=True, after=1)
    status_badge(doc, status)
    tbl = table(doc,
        ["Aspect", "Detail"],
        [
            ["Trigger", trigger],
            ["What the AI shows", ai_shows],
        ],
        header_fill=HDR_FILL)
    para(doc, "CSR options:", size=10, bold=True, before=3, after=1)
    bullets(doc, options)
    para(doc, "How to test:", size=10, bold=True, before=2, after=1)
    for i, stp in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(stp).font.size = Pt(10)

# ── 4. Test matrix ────────────────────────────────────────────────────────────
heading(doc, "4. Quick Test Matrix")
para(doc, "One-line trigger for each CSR condition:")
table(doc,
    ["#", "Condition", "Layer", "Minimal trigger"],
    [
        ["3.1", "Missing/unresolved buyer", "Intake", "Blank/unknown buyer email"],
        ["3.2", "Partial ship-to", "Intake", "CSR-Approval-PO (name-only ship-to)"],
        ["3.3", "Missing SKU", "Intake", "CSR-Approval-PO line 3 (desc + qty only)"],
        ["3.4", "Wrong/invalid SKU", "Intake", "CSR-Approval-PO line 2 (PN-DRAIN-STD)"],
        ["3.5", "Zero quantity", "Intake", "CSR-Approval-PO line 4 (Qty = 0)"],
        ["3.6", "Duplicate PO", "Intake", "Resubmit PO-2026-30001"],
        ["3.7", "Obsolete SKU", "Product Match", "CSR-Approval-PO line 1 (SKU-CTG-1000)"],
        ["3.8", "Pricing exception", "Pricing and Promo", "SKU-CTG-4520 qty 2000 (CUST-1001)"],
        ["3.9", "Credit hold", "Credit", "Customer CUST-7000 / CUST-1002"],
        ["3.10", "Inventory shortage", "Inventory Checks", "SKU-SHS-7700 qty 10"],
        ["3.11", "ZIP not serviceable", "Logistics", "Ship-to ZIP 99950"],
        ["3.12", "Approval routing", "Approval", "Order > $250K, or buyer Linda Park"],
        ["3.13", "Strategic customer", "Customer Validation", "CUST-1001 (Strategic tier)"],
    ],
    header_fill=HDR_FILL)

# ── 5. Two-PO demo strategy ───────────────────────────────────────────────────
heading(doc, "5. Recommended Demo Strategy (two PO files)")
para(doc, "For a clean client demonstration, two PO files cover the full story:")
bullets(doc, [
    ("Happy-Flow-PO (Happy path)", "PO-2026-30001 - flows straight through all nine "
     "layers with no CSR interruption; shows auto-approval, pricing, credit pass, "
     "inventory, logistics and the confirmed order."),
    ("CSR-Approval-PO (Exceptions)", "PO-2026-30002 - exercises every interactive CSR "
     "gate in one run: unknown buyer, obsolete/wrong/missing SKU, zero qty, UOM, "
     "partial ship-to, then the Pricing, Credit, Inventory and Logistics gates."),
])
para(doc, "CSR-Approval-PO now carries the Pricing, Credit, Inventory and Logistics "
          "triggers as well, so all interactive CSR gates can be shown from that one "
          "file. The individual triggers in section 4 remain available for ad-hoc "
          "demonstration of a single gate in isolation.",
     size=10)

# ── 6. Sign-off ───────────────────────────────────────────────────────────────
heading(doc, "6. Review Sign-Off")
para(doc, "Please confirm the following before implementation proceeds:")
bullets(doc, [
    "The nine decision-layer headings and the mapping in section 2 are correct.",
    "Folding Buyer Authorization into Customer Validation and Compliance into "
    "Product Match is acceptable (extra stages hidden as sub-checks).",
    "The Pricing and Promo layer should become an interactive CSR gate (3.8).",
    "The Product Match substitution buttons should be Approve / Modify / Escalate "
    "(no Reject) per 3.7.",
    "Customer Tier, Payment Terms and Distributor classification should appear in "
    "the Resolved Account Hierarchy.",
])

foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(12)
fr = foot.add_run("EY")
fr.bold = True
fr.font.size = Pt(14)
fr.font.color.rgb = EY_DARK

out = "CSR_Approval_Conditions_and_Test_Guide.docx"
doc.save(out)
print("Saved:", out)
