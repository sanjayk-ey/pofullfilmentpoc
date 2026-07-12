"""
create_demo_start_guide.py
Generates docs/Demo_Start_Guide.docx — a very simple, step-by-step guide for the
business team to start and run the demo today. Plain language, no jargon.

Run:  python create_demo_start_guide.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(__file__), "docs")
os.makedirs(OUT_DIR, exist_ok=True)

NAVY = RGBColor(0x1F, 0x2A, 0x44)
BLUE = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x1E, 0x7D, 0x44)
AMBER = RGBColor(0xB4, 0x53, 0x09)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("PO Fulfilment Order Assistant")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("Demo Start Guide — Simple Steps")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared by: Sanjay Kumar Kesarvani    |    Date: 30 June 2026").font.size = Pt(10)

    doc.add_paragraph()

    def h(text, color=NAVY, size=14):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p

    def step(num, text):
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)
        return p

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
        return p

    # ---- What this is ----
    h("What is this?")
    doc.add_paragraph(
        "This is a smart assistant that reads a Purchase Order (PO) and checks it from start to "
        "finish — just like an experienced order desk would. It pulls out the details, checks the "
        "customer, prices, budget, credit, stock and delivery, and finally creates the order. "
        "If something is wrong, it stops, explains the problem in plain words, and points it to "
        "the right team."
    )

    # ---- Before the demo ----
    h("Before the demo (one-time setup)")
    doc.add_paragraph(
        "Your engineer normally does this part. If you need to do it yourself, follow these steps "
        "on the demo laptop:"
    )
    step(1, "Open the Start menu, type 'cmd', and press Enter to open a black command window.")
    step(2, "Type this and press Enter:  cd C:\\projects\\po-fullfiment-poc")
    step(3, "Turn on the work area. Type:  venv\\Scripts\\activate   (you will see (venv) appear).")
    step(4, "Install the parts (first time only). Type:  python -m pip install -r requirements.txt")
    doc.add_paragraph(
        "That's it for setup. You only need to do this once on a new laptop.", style="Intense Quote"
    )

    # ---- Start the demo ----
    h("How to START the demo (do this today)", GREEN)
    step(1, "Open the command window (Start menu -> type 'cmd' -> Enter).")
    step(2, "Go to the project folder. Type:  cd C:\\projects\\po-fullfiment-poc")
    step(3, "Turn on the work area. Type:  venv\\Scripts\\activate")
    step(4, "Start the app. Type:  python -m streamlit run app.py")
    step(5, "Your web browser opens automatically and shows a chat screen. "
            "If it does not open, type this address in your browser:  http://localhost:8501")
    doc.add_paragraph(
        "You are ready. Leave the black command window open in the background — do not close it "
        "while you demo.", style="Intense Quote"
    )

    # ---- Two ways to send a PO ----
    h("Two ways to give a PO to the agent")
    p = doc.add_paragraph()
    p.add_run("A) Paste text:  ").bold = True
    p.add_run("Open a PO .txt file, select all the text, copy it, paste it into the chat box at "
              "the bottom of the screen, and press Enter.")
    p = doc.add_paragraph()
    p.add_run("B) Upload Excel:  ").bold = True
    p.add_run("Click the + (plus) button near the chat box, then choose a PO Excel (.xlsx) file. "
              "Only Excel files are accepted. Ready-to-use Excel copies of the two demo POs "
              "live at demo/Happy-Flow-PO.xlsx and demo/CSR-Approval-PO.xlsx.")

    # ---- What you will see ----
    h("What you will see on screen")
    bullet("The agent shows each step as it works, one after another.")
    bullet("GREEN / a tick means the step is fine and it moves on.")
    bullet("An AMBER / RED box means it found a problem. It explains the problem in simple words "
           "and says who should handle it. This is the correct behaviour for the 'problem' demos.")
    bullet("For approvals or pricing problems, it shows the email being sent to the right "
           "manager, and the process pauses until that manager responds.")
    bullet("When everything passes, it creates the order and shows confirmation messages.")

    # ---- Suggested running order ----
    h("Suggested demo flow (only TWO PO files)", BLUE)
    doc.add_paragraph(
        "The whole demo now uses just two purchase-order files. Run them in this order:")
    flow = [
        ("1. The happy path", "demo/Happy-Flow-PO.txt  (or .xlsx)",
         "A clean order that contains only the mandatory fields. Show the agent "
         "reading it, resolving the customer and buyer from the company name and "
         "email, and running every check straight through to a created order with "
         "no human input needed. Every optional header field (contact person, "
         "contract reference, delivery instructions) is backfilled from master "
         "data — the PO card marks these with a small 'from master data' badge. "
         "During pricing the agent internally calculates BOTH the freight / "
         "shipping charge and the state sales tax and displays them in the "
         "'Order totals' block."),
        ("2. The interactive CSR-approval PO", "demo/CSR-Approval-PO.txt  (or .xlsx)",
         "One order that deliberately contains problems, so you can show the agent "
         "thinking and asking the CSR to decide at every human-in-the-loop gate. In "
         "sequence it will: (a) ask which registered buyer to use when the email is "
         "unknown; (b) recommend a substitute for a discontinued product and ask for "
         "approval; (c) identify a product from its description when the customer used "
         "their own code; (d) identify a product when the SKU is missing entirely; "
         "(e) ask for a quantity when a line has zero; (f) convert a non-standard unit "
         "of measure (cases to each) and show the maths; (g) confirm a partial ship-to "
         "(a location name only) against the address book; then the decision-layer "
         "gates — (h) a Pricing & Promo discount exception, (i) a Credit hold, "
         "(j) an Inventory shortage, and (k) an unserviceable delivery ZIP. For each "
         "one the CSR can Approve, Reject, or Escalate — or type a correction. Approve "
         "each in turn and the order completes."),
    ]
    for title, fpath, why in flow:
        p = doc.add_paragraph()
        p.add_run(title + "  ").bold = True
        p.add_run("(" + fpath + ")").italic = True
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(why)

    doc.add_paragraph(
        "Tip: the agent works deliberately at a steady pace and shows what it is "
        "checking internally (like an assistant thinking) so the audience can follow "
        "each decision.")

    doc.add_paragraph(
        "For the full list of every test and its expected result, see the Excel file "
        "'Unit_Test_Plan_Interactive_Demo.xlsx' in the docs folder.", style="Intense Quote"
    )

    # ---- If something goes wrong ----
    h("If something goes wrong", AMBER)
    bullet("The screen looks stuck or shows an error: refresh the browser page (press F5).")
    bullet("The browser tab closed: open it again at  http://localhost:8501")
    bullet("The app stopped completely: go back to the black command window and run "
           "'python -m streamlit run app.py' again.")
    bullet("An upload is rejected: make sure the file is a .xlsx Excel file (not PDF or Word).")

    # ---- How to stop ----
    h("How to STOP the demo")
    step(1, "Close the browser tab.")
    step(2, "Go to the black command window and press the Ctrl key and C together to stop the app.")

    out = os.path.join(OUT_DIR, "Demo_Start_Guide.docx")
    doc.save(out)
    print("Created:", out)


if __name__ == "__main__":
    main()
