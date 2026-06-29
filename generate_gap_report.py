"""
Generates Sample_Data_Gap_Report.docx — a formal gap analysis document
for the business team based on review of Sample_data.xlsx against all
12 user story acceptance criteria.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import datetime

# ── helpers ────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def _set_cell_bg(cell, hex_color):
    """Apply background fill to a table cell via raw XML."""
    from lxml import etree
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove any existing shd element
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd_elem = etree.SubElement(tcPr, qn("w:shd"))
    shd_elem.set(qn("w:val"), "clear")
    shd_elem.set(qn("w:color"), "auto")
    shd_elem.set(qn("w:fill"), hex_color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(cell, "2E4057")
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        # alternate row shading
        if ri % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, "F5F7FA")
    if col_widths:
        for ri2, row in enumerate(table.rows):
            for ci2, cell in enumerate(row.cells):
                if ci2 < len(col_widths):
                    cell.width = Cm(col_widths[ci2])
    return table

# ── gap data ───────────────────────────────────────────────────────

GAPS = [
    # id, tab, story, severity, title, detail
    ("G-01","po_intake_requests","US-01","CRITICAL",
     "Missing SKU, Quantity, and UOM columns",
     "AC-01 requires extracting SKU, quantity, and unit of measure from every PO. "
     "The po_intake_requests tab captures customer, buyer, PO number, ship-to, date, "
     "and delivery instructions — but has NO columns for ordered items. Without these, "
     "the entire downstream pipeline (product match, pricing, inventory, logistics) "
     "cannot be fed. Add: sku, quantity, uom (one row per line item, or a separate "
     "po_intake_lines child table)."),

    ("G-02","po_intake_requests","US-01","HIGH",
     "Missing contract_reference column",
     "AC-01 explicitly requires 'contract reference' extraction. No column exists. "
     "Add: contract_reference_raw. The pricing engine (US-06) uses this to resolve "
     "the correct contract price."),

    ("G-03","po_intake_requests","US-01","MEDIUM",
     "Missing cost_center column in intake",
     "US-03 AC-03 validates cost center assignment from the PO. The intake record "
     "should capture cost_center_raw so buyer authorization can validate it."),

    ("G-04","po_intake_requests","US-01","MEDIUM",
     "ship_to_raw is free-text with no ZIP code",
     "AC-01 requires 'ship-to ZIP' extraction. ship_to_raw is unstructured text "
     "('Great Lakes Manufacturing receiving dock, Chicago, IL'). Either add a "
     "ship_to_zip column, or ensure every ship_to_raw value includes the ZIP code."),

    ("G-05","po_intake_requests + orders","US-01","HIGH",
     "Duplicate PO (INT-US-9006) still creates a full order",
     "AC-03 requires blocking duplicate POs. INT-US-9006 reuses PO number USPO-26-4201 "
     "and has status='duplicate', yet ORD-US-60006 is still created for it. "
     "Orders for duplicate or exception intakes should NOT exist."),

    ("G-06","accounts","US-02","CRITICAL",
     "Hierarchy is flat — missing Global Parent, Regional Division, and Local Branch levels",
     "AC-01 requires identifying 'global parent company, regional division, local branch, "
     "and specific ship-to location'. The accounts tab references parent_account_id values "
     "(US-GRP-100/200/300) but these group records are not defined as rows. The hierarchy "
     "is only 2 levels deep (Group -> Plant). Add: Global Parent rows, Regional Division "
     "rows, and Local Branch rows, each with their own account_id and hierarchy_level."),

    ("G-07","accounts","US-02","HIGH",
     "No hierarchy_rules / fulfillment_rules table",
     "AC-03 says 'apply the correct rules from the most specific eligible hierarchy level' "
     "for pricing, product visibility, budget, approval routing, and fulfillment. No table "
     "defines these per-level rules. Need a hierarchy_rules table with: hierarchy_level, "
     "account_id, preferred_warehouse, split_shipment_allowed, backorder_tolerance_days, "
     "pricing_tier, approval_routing, fulfillment_rule."),

    ("G-08","ship_to_locations","US-02","MEDIUM",
     "No invalid ship-to or hierarchy mismatch test scenario",
     "AC-02 requires escalation when ship-to is not associated with the account. All 20 "
     "ship-to locations map cleanly 1:1 with their account. Need at least one PO where "
     "the ship_to belongs to a different account to trigger the INVALID_SHIP_TO exception."),

    ("G-09","buyers","US-03","HIGH",
     "Missing product_visibility_rules table",
     "AC-02 requires validating 'product visibility rules' per branch/cost center. No "
     "product_visibility or buyer_product_access table exists. Add a table with: rule_id, "
     "scope (buyer_id/account_id/cost_center_id), product_family/sku, visibility_status "
     "(VISIBLE/RESTRICTED/HIDDEN), and reason."),

    ("G-10","buyers","US-03","MEDIUM",
     "Missing buyer authority / order limit columns",
     "AC-01 requires validating 'order limits and purchasing rights'. buyers.approval_threshold_amount "
     "exists but missing: max_order_limit, can_self_approve flag, allowed_product_families, "
     "allowed_branches."),

    ("G-11","approvers","US-03","MEDIUM",
     "Approvers have no account_id — cannot link to hierarchy",
     "approvers tab has no account_id column. The buyers.approver_id FK exists but there is "
     "no approval_matrix defining escalation chains per hierarchy level. US-07 AC-02 needs "
     "this to route approval to the correct regional or corporate approver."),

    ("G-12","product_master","US-04","HIGH",
     "Missing configurable product attributes",
     "AC-01 requires matching 'size, material, grade, compatibility, region, and configuration "
     "options'. product_master only has sku, description, family, status, uom, pack_size. "
     "Missing: material, grade, compatibility_rules, configuration_attributes, region_restriction. "
     "Add a product_attributes child table."),

    ("G-13","product_master","US-04","HIGH",
     "Missing substitution_rules table",
     "AC-04 requires 'approved substitute product' with 'compatibility rationale, price impact, "
     "and availability impact'. product_master has superseded_by_sku but no substitution_rules "
     "table with: original_sku, substitute_sku, rationale, price_delta_pct, availability_note, "
     "requires_csr_approval."),

    ("G-14","uom_conversions","US-04","MEDIUM",
     "18 of 20 UOM conversions are trivial identity (1:1, same UOM)",
     "Only FCT-US-2020 has a real conversion (lb -> case, factor=25). To demonstrate AC-02 "
     "(UOM conversion and validation logic), add more realistic conversions such as: "
     "box->each, gallon->litre, kg->lb, pallet->each."),

    ("G-15","compliance_rules","US-05","HIGH",
     "No test PO ordering a restricted product to a prohibited region",
     "AC-03 requires blocking restricted product-region combinations. RULE-US-9018 has "
     "ADP-US-9091 eligible=False in US-AZ, but no PO in po_intake_requests orders this "
     "SKU to Arizona. Add a PO that triggers this COMPLIANCE_RESTRICTION exception."),

    ("G-16","compliance_documents","US-05","MEDIUM",
     "No missing compliance document scenario",
     "AC-02 requires attaching SDS/compliance docs and logging when they are missing. All "
     "20 SKUs have a matching compliance document. Add at least one SKU where the required "
     "doc is absent or expired to trigger the MISSING_SDS exception."),

    ("G-17","pricing_contracts","US-06","HIGH",
     "Missing margin_policy / pricing exception routing table",
     "AC-05 requires routing pricing exceptions to 'the appropriate pricing approver'. "
     "Contracts have margin_floor_pct but there is no margin_policy or pricing_exception "
     "table defining discount ceiling, margin threshold, and approver_role for escalation."),

    ("G-18","pricing_contracts","US-06","MEDIUM",
     "Expired contract (CON-US-3012) not used to test date-bound pricing fallback",
     "AC-02 requires fallback when contract is expired. CON-US-3012 expires 2026-03-31 and "
     "the corresponding PO date is 2026-07-06 (past expiry), but orders.order_status does "
     "not show a pricing exception for ORD-US-60012. Explicitly wire this as a scenario."),

    ("G-19","price_waterfall_lines","US-06","MEDIUM",
     "Only 1 waterfall row per order line — full multi-layer breakdown missing",
     "AC-01 requires showing: list price -> contract -> volume tier -> promo -> surcharge -> "
     "final. Each order line has exactly 1 waterfall row. Add multiple sequential rows showing "
     "each pricing layer."),

    ("G-20","budgets","US-07","HIGH",
     "Missing approval_matrix table",
     "AC-02 requires an 'account hierarchy approval matrix' to route to the correct regional "
     "or corporate approver. No approval_matrix or approval_routing table exists. Add: "
     "matrix_id, account_id, hierarchy_level, spend_threshold_min, spend_threshold_max, "
     "approver_id, escalation_approver_id, sla_hours."),

    ("G-21","budgets","US-07","MEDIUM",
     "No order value actually exceeds remaining budget",
     "AC-03 requires exception when budget is unavailable. BUD-US-11004 has only $5000 "
     "remaining but ORD-US-60004 is only $944. Need a scenario where order total > "
     "remaining budget to trigger BUDGET_EXCEEDED."),

    ("G-22","approval_requests","US-07","MEDIUM",
     "Inconsistent approval outcome data",
     "ARQ-US-14013 status='pending' but has a decision_timestamp (it should be null). "
     "ARQ-US-14017 status='pending' but has a decision_timestamp. AC-04 requires capturing "
     "decision timestamp only when a decision is made."),

    ("G-23","credit_profiles","US-08","MEDIUM",
     "Missing payment_history and fraud_signals columns",
     "AC-01 requires checking 'payment risk'. Systems/Data Needed lists 'payment history, "
     "fraud/risk signals'. The credit_profiles tab has risk_category and credit_status but "
     "no: payment_history_score, last_payment_date, fraud_flag, overdue_invoice_count."),

    ("G-24","invoice_aging","US-08","MEDIUM",
     "Several invoices have invoice_amount = $0",
     "INV-US-12001, INV-US-12006, INV-US-12011, INV-US-12016 have invoice_amount=0 with "
     "status='open'. A $0 open invoice has no meaning. Replace with realistic invoice amounts."),

    ("G-25","inventory_balances","US-09","CRITICAL",
     "Only 1 warehouse per SKU — multi-DC sourcing, split shipment, and allocation cannot be tested",
     "AC-01 requires checking 'inventory across plant stock, DC inventory, in-transit stock, "
     "and ATP'. Every SKU appears at exactly 1 warehouse. To test split shipment (AC-02/AC-04), "
     "alternate warehouse sourcing, and allocation priority, each SKU must have stock rows at "
     "multiple warehouses with different quantities."),

    ("G-26","inventory_balances","US-09","HIGH",
     "Missing in_transit_qty and atp_qty columns",
     "AC-01 explicitly requires 'in-transit stock' and 'available-to-promise supply'. The tab "
     "has quantity_available and next_available_date but no in_transit_qty or atp_qty columns. "
     "These are separate inventory concepts that the ATP engine must use."),

    ("G-27","inventory_balances + ship_to_locations","US-09","HIGH",
     "Missing allocation_rules table",
     "AC-03 requires 'customer-specific fulfillment rules: preferred warehouse, restricted DC, "
     "split-shipment preference, backorder tolerance'. AC-04 requires 'allocation priority based "
     "on customer tier, contract commitment, order urgency'. No allocation_rules table exists. "
     "ship_to_locations has preferred_warehouse and split_shipment_allowed but missing: "
     "restricted_warehouse_ids, min_order_qty, delivery_sla_days, allocation_priority."),

    ("G-28","carrier_serviceability","US-10","HIGH",
     "Missing SLA rules and delivery calendar table",
     "AC-01 requires 'delivery SLA, delivery calendar, and transit time'. The tab has "
     "transit_days and freight_cost but no: sla_max_days, delivery_calendar_id, working_days, "
     "holiday_blackout_dates, delivery_window."),

    ("G-29","carrier_serviceability","US-10","MEDIUM",
     "Only 1 carrier per warehouse — no carrier comparison possible",
     "AC-02 and AC-04 require evaluating multiple carriers for optimal freight cost and SLA. "
     "Each warehouse has exactly 1 carrier. Add 2-3 carrier options per warehouse-ZIP "
     "combination to enable carrier comparison."),

    ("G-30","exceptions","US-11","HIGH",
     "Missing exception_governance / SLA config table",
     "AC-03 requires 'exception SLA threshold' and auto-escalation to the next role if "
     "unresolved. No exception_governance or sla_config table exists. Add: exception_type, "
     "sla_hours, escalation_role, auto_escalate_flag, notification_template."),

    ("G-31","exceptions","US-11","MEDIUM",
     "Exception assigned_role does not match exception_type",
     "EXC-US-15006 type='pricing' is assigned to 'Compliance Reviewer'. EXC-US-15004 "
     "type='credit' is assigned to 'Pricing Desk'. EXC-US-15010 type='compliance' is "
     "assigned to 'Pricing Desk'. Each exception type must route to the correct functional "
     "role: pricing->Pricing Desk, credit->Finance, compliance->Compliance Reviewer."),

    ("G-32","downstream_execution_events","US-12","HIGH",
     "Only 1 downstream event per order — should be 5 (ERP, OMS, WMS, TMS, Notification)",
     "AC-01 requires creating ERP sales order + OMS request + WMS pick ticket + TMS shipment "
     "+ customer confirmation for each completed order. Each order currently has 1 event. "
     "A fully executed order needs all 5 events as separate rows."),

    ("G-33","downstream_execution_events","US-12","MEDIUM",
     "Customer confirmation content columns missing",
     "AC-02 requires the confirmation to include: confirmed order number, PO number, final "
     "price, contract reference, quantity, UOM, fulfillment source, ETA, tracking number, "
     "payment terms, and compliance document references. None of these are columns in the "
     "execution events tab."),

    ("G-34","agent_audit_log","US-12","MEDIUM",
     "Audit log only covers 9 of 12 orchestration steps, with 1 entry per order",
     "AC-03 requires capturing every automated decision. The log cycles through 9 steps "
     "(intake, account_validation, uom_conversion, inventory_check, pricing_waterfall, "
     "approval_gate, credit_check, compliance_check, execution). Missing: buyer_authorization, "
     "product_match, logistics_validation. Also each order should have multiple entries "
     "(one per step), not a single row."),

    ("G-35","orders + po_intake_requests","Cross-cutting","HIGH",
     "Orders created for exception and duplicate intakes — data is logically inconsistent",
     "INT-US-9004 (confidence=0.62, status=exception) and INT-US-9006 (status=duplicate) "
     "both have corresponding orders in the orders tab. Intakes with status 'exception' or "
     "'duplicate' should NOT progress to order creation. Remove ORD-US-60004 and ORD-US-60006 "
     "or correct the intake statuses."),

    ("G-36","Multiple","Cross-cutting","HIGH",
     "All 20 orders are single-line — no multi-SKU POs",
     "Real B2B POs typically contain multiple line items. Every order in the dataset has "
     "exactly 1 order_line with 1 SKU. Add at least 3 orders with 2-4 line items each to "
     "test: multi-line pricing, partial fulfillment, split shipment across lines, and "
     "line-level exception handling."),

    ("G-37","po_intake_requests","US-01","MEDIUM",
     "Low-confidence intake has no missing_fields column",
     "AC-02 requires displaying 'which mandatory fields are missing' to the CSR. INT-US-9004 "
     "has confidence=0.62 and status=exception, but no column identifies which specific fields "
     "could not be extracted. Add a missing_fields column (comma-separated list)."),
]

MISSING_TABLES = [
    ("po_intake_lines (child table)", "US-01",
     "One row per line item in the PO: sku, quantity, uom, unit_price, line_total"),
    ("account_hierarchy (full 4-level)", "US-02",
     "Global Parent, Regional Division, Local Branch rows with hierarchy_level and parent_account_id"),
    ("hierarchy_rules", "US-02/09/10",
     "Per-level rules: preferred_warehouse, split_shipment, backorder_tolerance, pricing_tier, fulfillment_rule"),
    ("product_visibility_rules", "US-03",
     "SKU/family visibility per buyer role, account, cost center: VISIBLE / RESTRICTED / HIDDEN"),
    ("buyer_authority", "US-03",
     "max_order_limit, can_self_approve, allowed_product_families, allowed_branches"),
    ("product_attributes", "US-04",
     "material, grade, compatibility_rules, configuration_options, region_restriction per SKU"),
    ("substitution_rules", "US-04",
     "original_sku, substitute_sku, rationale, price_delta_pct, requires_csr_approval"),
    ("margin_policy", "US-06",
     "discount_ceiling_pct, margin_floor_pct, approver_role for pricing exception routing"),
    ("approval_matrix", "US-07",
     "account_id, hierarchy_level, spend_threshold_min/max, approver_id, escalation_approver_id, sla_hours"),
    ("allocation_rules", "US-09",
     "restricted_warehouse_ids, min_order_qty, delivery_sla_days, allocation_priority per account/ship-to"),
    ("sla_rules / delivery_calendar", "US-10",
     "max_delivery_days, working_days, holiday_blackout_dates, delivery_window per carrier/region"),
    ("exception_governance", "US-11",
     "exception_type, sla_hours, escalation_role, auto_escalate_flag, notification_template"),
]

# ── build document ──────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Title
t = doc.add_heading("Sample Data Gap Analysis Report", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Prepared by: Sanjay Kumar Kesarvani     |     Date: {datetime.date.today().strftime('%d %B %Y')}\n").font.size = Pt(10)
meta.add_run("Review of: Sample_data.xlsx  |  Against: US-01 to US-12 Acceptance Criteria").font.size = Pt(10)

doc.add_paragraph()

# ── 1. Executive Summary ───────────────────────────────────────────
heading(doc, "1. Executive Summary", 1)
para(doc,
     "A thorough review of the Sample_data.xlsx file provided by the business team was conducted "
     "against all 12 user story acceptance criteria, exception types, validation checks, and "
     "data completeness requirements for the PO Fulfillment AI Agent POC. "
     "The file contains 27 data tabs with 20 sample rows each (4,221 rows in the consolidated view). "
     "This review identified 37 gaps and 12 missing tables/sheets that must be addressed before "
     "the sample data can be used to test the AI orchestration pipeline.")
doc.add_paragraph()

# Summary stats table
heading(doc, "1.1 Severity Summary", 2)
sev_counts = {}
for g in GAPS:
    sev_counts[g[3]] = sev_counts.get(g[3], 0) + 1

add_table(doc,
    ["Severity", "Count", "Impact"],
    [
        ["CRITICAL", str(sev_counts.get("CRITICAL", 0)),
         "Entire pipeline stage cannot function — fix before any testing"],
        ["HIGH",     str(sev_counts.get("HIGH",     0)),
         "Key acceptance criteria cannot be demonstrated"],
        ["MEDIUM",   str(sev_counts.get("MEDIUM",   0)),
         "Specific exception scenarios or edge cases will fail"],
        ["LOW",      str(sev_counts.get("LOW",      0)),
         "Minor completeness gaps — low demo risk"],
        ["Missing Tables", str(len(MISSING_TABLES)),
         "Entire sheets absent from the file"],
    ],
    col_widths=[3, 2, 12]
)
doc.add_paragraph()

# ── 2. Missing Tables ──────────────────────────────────────────────
heading(doc, "2. Missing Tables / Sheets (12)", 1)
para(doc,
     "The following tables are required by the user story acceptance criteria and "
     "'Systems / Data Needed' columns but have NO corresponding sheet in Sample_data.xlsx. "
     "These must be added for the AI pipeline to function end-to-end.")
doc.add_paragraph()

add_table(doc,
    ["Table Name", "Story", "Required Columns / Purpose"],
    [[t[0], t[1], t[2]] for t in MISSING_TABLES],
    col_widths=[4.5, 2.5, 10]
)
doc.add_paragraph()

# ── 3. Detailed Gaps by Severity ──────────────────────────────────
heading(doc, "3. Detailed Gaps by Severity", 1)

for sev in ["CRITICAL", "HIGH", "MEDIUM"]:
    sev_gaps = [g for g in GAPS if g[3] == sev]
    if not sev_gaps:
        continue

    heading(doc, f"3.{'1' if sev=='CRITICAL' else '2' if sev=='HIGH' else '3'}.  {sev} Gaps  ({len(sev_gaps)})", 2)

    rows = []
    for g in sev_gaps:
        rows.append([g[0], g[4], f"Tab: {g[1]}\nStory: {g[2]}\n\n{g[5]}"])

    add_table(doc,
        ["ID", "Issue Title", "Detail"],
        rows,
        col_widths=[1.2, 4.5, 11.3]
    )
    doc.add_paragraph()

# ── 4. Data Consistency Quick-Fix List ────────────────────────────
heading(doc, "4. Data Consistency Quick-Fix List", 1)
para(doc,
     "The following corrections are needed in the EXISTING tabs (no new sheets required):")
doc.add_paragraph()

fixes = [
    ("po_intake_requests", "Add columns: sku, quantity, uom, contract_reference, cost_center_raw, ship_to_zip, missing_fields"),
    ("accounts", "Add rows for Global Parent (US-GRP-100/200/300), Regional Division, and Local Branch with correct hierarchy_level values"),
    ("approvers", "Add account_id column to link each approver to an account or hierarchy level"),
    ("inventory_balances", "Add multiple warehouse rows per SKU; add in_transit_qty and atp_qty columns"),
    ("orders", "Remove ORD-US-60004 and ORD-US-60006 (created for exception/duplicate intakes)"),
    ("invoice_aging", "Replace $0 invoice amounts (INV-US-12001/06/11/16) with realistic values"),
    ("exceptions", "Fix assigned_role: pricing->Pricing Desk, credit->Finance, compliance->Compliance Reviewer"),
    ("approval_requests", "Remove decision_timestamp from ARQ-US-14013 and ARQ-US-14017 (status=pending)"),
    ("agent_audit_log", "Add buyer_authorization, product_match, logistics_validation step types; add multiple entries per order"),
    ("order_lines", "Add 3+ orders with multiple SKU lines (2-4 per order)"),
    ("downstream_execution_events", "Add all 5 events (ERP, OMS, WMS, TMS, Notification) per completed order"),
    ("price_waterfall_lines", "Add sequential layer rows per order line (list, contract, volume, promo, surcharge, final)"),
    ("uom_conversions", "Replace trivial 1:1 identity conversions with realistic multi-unit conversions"),
    ("carrier_serviceability", "Add 2-3 carrier options per warehouse-ZIP pair to enable carrier comparison"),
]

add_table(doc,
    ["Tab", "Correction Required"],
    fixes,
    col_widths=[4.5, 12.5]
)
doc.add_paragraph()

# ── 5. Recommended Test Scenarios Missing ─────────────────────────
heading(doc, "5. Required Test Scenarios Not Present in Data", 1)
para(doc,
     "The following specific exception scenarios required by acceptance criteria are not "
     "represented in the current 20-order dataset. Business team must add at least one "
     "PO / data row to trigger each of these:")
doc.add_paragraph()

scenarios = [
    ("INVALID_SHIP_TO (US-02 AC-02)", "A PO where ship_to_location_id belongs to a different account than the buyer's account"),
    ("HIERARCHY_MISMATCH (US-02 AC-03)", "A PO where branch-level rule conflicts with parent-level rule to show cascade resolution"),
    ("RESTRICTED_PRODUCT (US-03 AC-02)", "A buyer ordering a product marked RESTRICTED for their cost center or branch"),
    ("COMPLIANCE_RESTRICTION (US-05 AC-03)", "A PO ordering ADP-US-9091 to Arizona (US-AZ) — eligible_flag=False in compliance_rules"),
    ("MISSING_SDS (US-05 AC-02)", "A PO for a hazardous SKU whose required compliance document is absent or expired"),
    ("PRICING_EXCEPTION (US-06 AC-05)", "A PO where the applied discount + promo breaches the margin_floor_pct"),
    ("BUDGET_EXCEEDED (US-07 AC-03)", "A PO where order total > budget_total - budget_used"),
    ("CREDIT_HOLD (US-08 AC-02)", "ACC-US-1005 already has credit_status='hold' — wire a PO for this account to trigger credit hold"),
    ("SPLIT_NOT_ALLOWED (US-09 AC-03)", "A PO where split_shipment_allowed=False but no single warehouse has full quantity"),
    ("INVENTORY_SHORTAGE (US-09 AC-02)", "A PO requesting more quantity than quantity_available at all warehouses combined"),
    ("ZIP_NOT_SERVICEABLE (US-10 AC-03)", "WH-PHX-01 to ZIP 850xx has serviceable_flag=False — wire a PO to trigger this"),
    ("EXECUTION_FAILURE (US-12 AC-04)", "A PO where ERP/OMS/WMS/TMS create fails — only ORD-US-60019 (TMS failed) is present, need more variety"),
]

add_table(doc,
    ["Scenario / Exception", "How to Create the Test Data"],
    [[s[0], s[1]] for s in scenarios],
    col_widths=[5, 12]
)
doc.add_paragraph()

# ── 6. Action Items ───────────────────────────────────────────────
heading(doc, "6. Recommended Next Steps for Business Team", 1)

steps = [
    "Add the 12 missing tables listed in Section 2, starting with po_intake_lines and approval_matrix (highest dependency).",
    "Fix the 3 CRITICAL gaps (G-01, G-06, G-25) immediately — these block the entire pipeline.",
    "Correct all 14 data consistency issues listed in Section 4.",
    "Add the 12 missing exception scenarios in Section 5 — at least one PO per exception type.",
    "Expand inventory_balances to have 2-3 warehouse entries per SKU to enable multi-DC sourcing.",
    "Expand orders to include at least 3 multi-line POs (2-4 SKUs each).",
    "Review exception routing in the exceptions tab and ensure assigned_role matches exception_type.",
    "Once corrections are made, share the updated file — we will re-validate and map to our POC mock data.",
]

for i, s in enumerate(steps):
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)

doc.add_paragraph()

# Footer note
note = doc.add_paragraph()
note.add_run(
    "Note: This analysis was performed against the POC scope covering US-01 to US-12. "
    "All integration points (ERP, OMS, WMS, TMS, SMTP) are mocked locally; real system "
    "connections are not required for the POC. However, the mock data must accurately reflect "
    "the business rules and exception patterns defined in the acceptance criteria."
).font.size = Pt(9)
note.runs[0].italic = True

# save
out_path = r"C:\projects\po-fullfiment-poc\docs\Sample_Data_Gap_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
