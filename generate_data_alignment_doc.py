"""
generate_data_alignment_doc.py
Generates docs/Data_Alignment_with_Business_Sample.docx

A reviewer-facing document that explains, field by field, how the POC mock data
and sample purchase orders were corrected and enriched using the business team's
reference workbook (Sample_data.xlsx) as the source of truth for schema and
relationships. Product names, units of measure and company entities were made to
look like real production data for a kitchen & bath fixtures manufacturer
(reference catalogue only; the client name is never used anywhere in the data).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.join(os.path.dirname(__file__), "docs")
os.makedirs(OUT_DIR, exist_ok=True)

NAVY = (30, 58, 95)
BLUE = (37, 99, 235)
GREY = (90, 90, 90)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _set_cell_bg(cell, hex_color):
    from lxml import etree
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(cell, "1E3A5F")
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
        if ri % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, "EEF3FA")
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ── Title block ─────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Mock & Sample Data Alignment with Business Reference Data")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(*NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("PO-to-Fulfillment Orchestration POC — Data Realism & Schema Alignment")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(*GREY)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Prepared by: Sanjay Kumar Kesarvani     |     Date: 24 June 2026")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(*GREY)
doc.add_paragraph()

# ── 1. Purpose ──────────────────────────────────────────────────────────────
heading(doc, "1. Purpose", level=1)
para(doc,
     "The business team shared a reference workbook (Sample_data.xlsx) describing "
     "the canonical data model for the order-fulfillment process. The team confirmed "
     "that the POC may continue to use its own locally generated mock data and sample "
     "purchase orders, but that this reference workbook should be used to (a) align "
     "the schema and the relationships between data sets, and (b) add any missing "
     "columns and values so that no unexpected checks or validations fail during the "
     "client demonstration.")
para(doc,
     "In addition, the data has been reworked so that it reads like genuine production "
     "data for a kitchen & bath fixtures business. The catalogue of products and units "
     "of measure was modelled on a real fixtures manufacturer's public product range "
     "(used only as a reference for realism). To keep the POC neutral, the client's "
     "name does not appear anywhere in the data, and words such as \u201Csample\u201D "
     "or \u201Cdemo\u201D have been removed from the data content itself.")

# ── 2. How the reference workbook was interpreted ───────────────────────────
heading(doc, "2. How the Reference Workbook Was Interpreted", level=1)
para(doc,
     "The reference workbook contains 27 inter-related tables. The \u201CSchema "
     "Details\u201D tab was read in full (212 column definitions) together with the "
     "\u201CKey / Relationship\u201D column, which identifies each primary key (PK) and "
     "foreign-key (FK) relationship. This established the entity-relationship map that "
     "the POC data now mirrors.")
para(doc, "Key relationship chains followed from the reference schema:", bold=True)
bullet(doc, "accounts (PK account_id) -> ship_to_locations, buyers, cost_centers, "
            "pricing_contracts, promotions, surcharges, budgets, credit_profiles, "
            "invoice_aging, orders.")
bullet(doc, "product_master (PK sku) -> uom_conversions, pricing_contracts, promotions, "
            "surcharges, inventory_balances, compliance_rules, compliance_documents, "
            "order_lines (and superseded_by_sku self-reference for obsolete items).")
bullet(doc, "warehouses (PK warehouse_id) -> ship_to_locations.preferred_warehouse_id, "
            "inventory_balances, carrier_serviceability, order_lines.fulfillment_warehouse_id.")
bullet(doc, "buyers.approver_id -> approvers; buyers.default_cost_center_id -> cost_centers; "
            "budgets -> accounts + cost_centers.")
para(doc,
     "The POC keeps these same relationships, but organises the data into ten "
     "domain workbooks (one per orchestration stage) instead of a single 27-tab "
     "file. The mapping between the two layouts is given in Section 3.")

# ── 3. Table mapping ────────────────────────────────────────────────────────
heading(doc, "3. Reference Table -> POC Workbook / Sheet Mapping", level=1)
add_table(doc,
    ["Reference table (27)", "POC workbook", "POC sheet(s)"],
    [
        ["po_intake_requests", "(runtime)", "Captured live by extractor.py from the PO"],
        ["accounts", "customer-master-data.xlsx", "Customer_Master, Account_Hierarchy"],
        ["ship_to_locations", "customer-master-data.xlsx", "Ship_To_Master"],
        ["buyers", "buyer-master-data.xlsx", "Buyer_Profiles, User_Permissions"],
        ["cost_centers", "buyer/budget-master-data.xlsx", "Cost_Centers"],
        ["approvers", "budget-master-data.xlsx", "Approval_Matrix"],
        ["product_master", "product-master-data.xlsx", "Product_Master, Product_Attributes"],
        ["uom_conversions", "product-master-data.xlsx", "UOM_Conversions"],
        ["pricing_contracts", "pricing-master-data.xlsx", "Contracts, Price_List"],
        ["volume_tiers", "pricing-master-data.xlsx", "Volume_Tiers"],
        ["promotions", "pricing-master-data.xlsx", "Promotions"],
        ["surcharges", "pricing-master-data.xlsx", "Surcharges"],
        ["warehouses", "logistics-master-data.xlsx", "Warehouse_Master"],
        ["inventory_balances", "inventory-master-data.xlsx", "Plant_Stock, DC_Stock, In_Transit, ATP"],
        ["carrier_serviceability", "logistics-master-data.xlsx", "Carrier_Coverage, Freight_Rating"],
        ["compliance_rules", "compliance-master-data.xlsx", "Compliance_Rules, Regional_Restrictions"],
        ["compliance_documents", "compliance-master-data.xlsx", "SDS_Repository"],
        ["budgets", "budget-master-data.xlsx", "Budget_Master, Cost_Centers"],
        ["credit_profiles", "credit-master-data.xlsx", "Credit_Master, Risk_Signals"],
        ["invoice_aging", "credit-master-data.xlsx", "Invoice_Aging, Payment_History"],
        ["orders / order_lines", "(runtime)", "Built by pipeline; ERP order id mocked at execution"],
        ["price_waterfall_lines", "(runtime)", "Produced by pricing_engine.py"],
        ["approval_requests", "(runtime)", "Produced by budget_approval.py + mock email"],
        ["exceptions", "exception-governance-master-data.xlsx", "Severity_Matrix, Role_Routing"],
        ["agent_audit_log", "(runtime)", "Each stage logs decision + rule + confidence"],
        ["downstream_execution_events", "execution-master-data.xlsx", "Integration_Endpoints (mocked)"],
    ],
    widths=[5.5, 5.5, 6.5])

# ── 4. Product catalogue adjustments ────────────────────────────────────────
heading(doc, "4. Product Catalogue Adjustments (realistic kitchen & bath range)", level=1)
para(doc,
     "The previous catalogue used generic industrial items (steel pipe, flanges, "
     "valves, pumps, gaskets, chemicals). These were replaced with a realistic "
     "kitchen & bath fixtures range. Internal SKU identifiers, list prices and "
     "quantities were preserved so that all pricing, budget, credit and inventory "
     "thresholds continue to behave exactly as validated, while descriptions, "
     "product families, materials and units of measure now reflect real fixtures.")
add_table(doc,
    ["SKU", "New description", "Family", "Base UOM", "List price"],
    [
        ["SKU-CTG-4520", "Ceramic Disc Faucet Cartridge, Single-Control", "CARTRIDGE", "EA", "$12.50"],
        ["SKU-DRN-3010", "Pop-Up Drain Assembly, Brushed Nickel", "DRAIN", "EA", "$45.00"],
        ["SKU-VLV-2201", "Pressure-Balancing Shower Valve, 1/2 in", "VALVE", "EA", "$88.00"],
        ["SKU-SHS-7700", "Digital Shower Interface System", "SHOWERSYS", "EA", "$1,450.00"],
        ["SKU-FIN-9100", "Enameled Cast Iron Touch-Up Finish, White", "FINISH", "GAL", "$9.75"],
        ["SKU-FIN-9200", "Cast Iron Touch-Up Aerosol, Almond (SDS pending)", "FINISH", "GAL", "$14.00"],
        ["SKU-SEL-1150", "Tank-to-Bowl Gasket Kit", "SEAL", "EA", "$6.25"],
        ["SKU-CTG-1000", "Legacy 2-Handle Faucet Cartridge (obsolete)", "CARTRIDGE", "EA", "$11.00"],
        ["SKU-VLV-2000", "Discontinued Thermostatic Shower Valve (inactive)", "VALVE", "EA", "$70.00"],
    ],
    widths=[3.2, 7.5, 3.0, 2.0, 2.3])
para(doc, "Units of measure now reflect how fixtures and fittings actually ship:", bold=True)
bullet(doc, "EA (each) for fixtures, cartridges, valves, drains, seals and the digital system.")
bullet(doc, "GAL for liquid touch-up finishes; CASE / BOX / PR conversions to EA added to "
            "UOM_Conversions for cartridges, seal kits and drains.")
bullet(doc, "An invalid request (e.g. ordering a cartridge in KG) still has no conversion "
            "path and correctly raises INVALID_UOM.")

# ── 5. Company / hierarchy / ship-to adjustments ────────────────────────────
heading(doc, "5. Account, Hierarchy and Ship-To Adjustments", level=1)
para(doc,
     "Placeholder corporate names (Acme / Globex / Initech) were replaced with "
     "realistic kitchen & bath distribution entities. Account IDs, branch codes and "
     "ship-to ZIP codes were preserved so the account-hierarchy and ship-to "
     "validation logic and all test scenarios remain intact.")
add_table(doc,
    ["Account ID", "New company name", "Branch", "Hierarchy parent"],
    [
        ["CUST-1001", "Great Lakes Plumbing Supply Co", "BR-GLP-MW", "Continental Building Products Group"],
        ["CUST-1002", "Eastern Kitchen & Bath Distributors", "BR-GLP-NE", "Continental Building Products Group"],
        ["CUST-2001", "Continental Canada Distribution", "BR-GLP-CA", "Continental Building Products Group"],
        ["CUST-5001", "Pacific Coast Bath & Kitchen", "BR-PCBK-WEST", "Western Supply Holdings"],
        ["CUST-7000", "Midtown Building Supply (x2 — duplicate test)", "BR-GLP-MW / NE", "Continental Building Products Group"],
    ],
    widths=[2.6, 7.0, 3.2, 5.0])
para(doc, "Ship-to locations keep their original ZIP codes (60639, 48201, 10001, 90001, "
          "99950, 90210) so the serviceability, hierarchy-mismatch and invalid-ship-to "
          "scenarios still resolve correctly; only the names and addresses were made "
          "realistic. Contract references were renamed from CONTRACT-ACME-* / "
          "CONTRACT-GLOBEX-* to CONTRACT-GLP-* / CONTRACT-PCBK-*.")

# ── 6. Columns / values added or confirmed per domain ───────────────────────
heading(doc, "6. Columns and Values Confirmed Against the Reference Schema", level=1)
para(doc,
     "Each domain workbook was checked against the corresponding reference table to "
     "ensure every column the validators rely on is present and populated, so the "
     "demo does not stop on a missing attribute:")
add_table(doc,
    ["Domain", "Reference fields confirmed / aligned"],
    [
        ["Accounts / hierarchy", "account_id, parent_account_id, hierarchy_level, account_status, customer_tier"],
        ["Ship-to", "preferred_warehouse_id, split_shipment_allowed, backorder_tolerance_days, postal_code"],
        ["Buyers", "role, email, approval_threshold_amount, approver_id, default_cost_center_id, buyer_status"],
        ["Product", "base_uom, pack_size/pack_unit, superseded_by_sku, hazardous_flag, requires_compliance_docs"],
        ["UOM", "from_uom, to_uom, conversion_factor, rounding_rule, material_change_threshold_pct"],
        ["Pricing", "list_price, contract_price, margin_floor_pct, volume tiers, promotions, surcharges"],
        ["Compliance", "country_region, eligible_flag, approval_required_flag, required_doc_type, SDS validity"],
        ["Budget / approval", "budget_total, budget_used, approver limits, urgency, justification"],
        ["Credit", "credit_limit, available_credit, payment_terms, risk_category, credit_status, days_overdue"],
        ["Inventory", "quantity_available, quantity_reserved, lead_time_days, next_available_date"],
        ["Logistics", "postal_code_prefix, carrier_name, serviceable_flag, transit_days, freight_cost"],
        ["Exceptions", "exception_type, severity, status, assigned_role, ai_recommendation"],
    ],
    widths=[4.0, 12.5])

# ── 7. Sample PO adjustments ────────────────────────────────────────────────
heading(doc, "7. Sample Purchase Order Adjustments", level=1)
bullet(doc, "All 30+ sample PO files (text and Excel) were regenerated with the new "
            "products, companies, ship-to names and realistic PO numbers (e.g. PO-2026-100xx).")
bullet(doc, "Embedded test annotations such as \u201CScenario:\u201D and \u201CExpected "
            "result:\u201D were removed so each file reads like a genuine purchase order.")
bullet(doc, "Words like \u201Csample\u201D and \u201Cdemo\u201D were removed from the PO "
            "content (titles, PO numbers, instructions).")
bullet(doc, "Each exception scenario still triggers exactly the intended exception "
            "(unauthorized buyer, restricted product, invalid UOM, obsolete SKU, "
            "compliance restriction, missing SDS, pricing breach, budget, approval, "
            "credit hold, inventory shortage, split-not-allowed, minimum order quantity, "
            "ZIP not serviceable, execution failure).")

# ── 8. Engine adjustment ────────────────────────────────────────────────────
heading(doc, "8. One Supporting Engine Adjustment", level=1)
para(doc,
     "Because real PO numbers now contain five-digit sequences (e.g. PO-2026-10012), "
     "the data-extraction logic was hardened so that an explicitly labelled "
     "\u201CShip-To ZIP\u201D value is always preferred over any incidental five-digit "
     "run in the document. This prevents a PO number from ever being mistaken for a ZIP "
     "code. No other validator logic was changed.")

# ── 9. Test results ─────────────────────────────────────────────────────────
heading(doc, "9. Verification — Full Regression", level=1)
para(doc,
     "After the data was corrected, the end-to-end regression harness "
     "(test_pipeline.py) was run across every scenario. Result: 27 / 27 cases pass "
     "\u2014 all happy paths flow cleanly through the full 12-stage pipeline, and "
     "every exception scenario raises the correct exception at the correct stage.")
add_table(doc,
    ["Category", "Scenarios", "Outcome"],
    [
        ["Happy path / autonomous", "US-01 (text, Excel, comprehensive x2), US-03, US-06, US-09 restricted-warehouse, US-11, US-12", "Clean pass"],
        ["Intake / account (US-01/02)", "missing fields, duplicate PO, unmatched / duplicate customer, invalid ship-to, hierarchy mismatch", "Correct exception"],
        ["Authorization / product (US-03/04)", "unauthorized buyer, restricted product, invalid cost center, obsolete SKU, invalid UOM, unknown SKU", "Correct exception"],
        ["Compliance / pricing (US-05/06)", "restricted region, missing SDS, pricing breach", "Correct exception"],
        ["Budget / credit (US-07/08)", "budget exceeded, approval required, credit hold", "Correct exception"],
        ["Inventory / logistics (US-09/10)", "inventory shortage, split-not-allowed, minimum order qty, ZIP not serviceable", "Correct exception"],
        ["Execution (US-12)", "execution failure (mock WMS outage)", "Correct exception"],
    ],
    widths=[5.0, 9.0, 2.5])

para(doc,
     "Note: the reference workbook also defines several reporting / audit tables "
     "(price_waterfall_lines, agent_audit_log, downstream_execution_events, orders, "
     "order_lines, approval_requests). In the POC these are produced at runtime by the "
     "orchestration engine rather than pre-loaded as master data, which is why they do "
     "not appear as static workbooks.", italic=True, color=GREY, size=9.5)

out = os.path.join(OUT_DIR, "Data_Alignment_with_Business_Sample.docx")
doc.save(out)
print("Created:", out)
